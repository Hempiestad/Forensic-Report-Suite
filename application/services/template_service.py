"""application/services/template_service.py — Full TemplateService implementation (Phase 5)."""
from __future__ import annotations

import re
from html.parser import HTMLParser
from typing import Any, Dict, List, Optional

from application.dtos.template_dto import (
    CreateTemplateDto,
    TemplateDto,
    TemplatePlaceholderDto,
    TemplateVersionDto,
    UpdateTemplateDto,
)
from application.interfaces.i_audit_service import IAuditService
from application.interfaces.i_clock import IClock
from application.interfaces.i_template_repository import ITemplateRepository
from application.interfaces.i_template_service import ITemplateService
from application.services._clock import DefaultClock
from domain.entities.template import Template
from domain.entities.template_placeholder import TemplatePlaceholder
from domain.enums.template_category import TemplateCategory
from domain.exceptions.domain_exceptions import DomainValidationError


# ---------------------------------------------------------------------------
# HTML well-formedness validator (no external deps)
# ---------------------------------------------------------------------------

class _HTMLValidator(HTMLParser):
    """Lightweight check: tracks open/close tags and reports mismatches."""

    _VOID = frozenset({
        "area", "base", "br", "col", "embed", "hr", "img", "input",
        "link", "meta", "param", "source", "track", "wbr",
    })

    def __init__(self) -> None:
        super().__init__()
        self._stack: list[str] = []
        self.errors: list[str] = []

    def handle_starttag(self, tag: str, attrs: list) -> None:
        if tag not in self._VOID:
            self._stack.append(tag)

    def handle_endtag(self, tag: str) -> None:
        if tag in self._VOID:
            return
        if self._stack and self._stack[-1] == tag:
            self._stack.pop()
        else:
            self.errors.append(f"Unexpected closing tag </{tag}>")

    def check(self, html: str) -> list[str]:
        self._stack.clear()
        self.errors.clear()
        try:
            self.feed(html)
        except Exception as exc:
            self.errors.append(f"Parse error: {exc}")
        for unclosed in self._stack:
            self.errors.append(f"Unclosed tag <{unclosed}>")
        return self.errors


# ---------------------------------------------------------------------------
# Mapper helpers
# ---------------------------------------------------------------------------

def _template_to_dto(tmpl: Template) -> TemplateDto:
    return TemplateDto(
        template_id=tmpl.id,
        name=tmpl.name,
        category=tmpl.category.value,
        description=tmpl.description or "",
        is_published=tmpl.is_published,
        is_system_template=tmpl.is_default,
        created_by=tmpl.created_by,
        created_at=tmpl.created_at,
        modified_at=tmpl.modified_at,
        modified_by=tmpl.modified_by,
        current_version=tmpl.version_number,
        usage_count=tmpl.usage_count,
        last_used_at=tmpl.last_used_at,
        tags=list(tmpl.tags),
        placeholders=[
            TemplatePlaceholderDto(
                name=ph.name,
                placeholder_type=ph.placeholder_type.value,
                description=ph.description,
                is_required=ph.is_required,
                default_value=ph.default_value,
                sample_value=ph.sample_value,
                token=ph.token,
            )
            for ph in tmpl.placeholders
        ],
    )


def _category_from_str(value: str) -> TemplateCategory:
    try:
        return TemplateCategory(value)
    except ValueError:
        # Fallback: match by name (e.g. "SWGDE_NIST" → TemplateCategory.SWGDE_NIST)
        for member in TemplateCategory:
            if member.name == value.upper():
                return member
        raise DomainValidationError("category", f"Unknown template category: {value!r}")


# ---------------------------------------------------------------------------
# Service
# ---------------------------------------------------------------------------

class TemplateService(ITemplateService):
    """Full TemplateService — implements all ITemplateService operations."""

    def __init__(
        self,
        template_repository: ITemplateRepository,
        audit_service: IAuditService,
        clock: Optional[IClock] = None,
    ) -> None:
        self._templates = template_repository
        self._audit = audit_service
        self._clock = clock or DefaultClock()
        self._next_id: int = 1  # Used only by in-memory repo; SQLite auto-increments

    # ------------------------------------------------------------------ #
    # Internal helpers                                                     #
    # ------------------------------------------------------------------ #

    def _next_template_id(self) -> int:
        """Return a next-available ID (heuristic for in-memory repo)."""
        all_ids = [t.id for t in self._templates.get_all()]
        return max(all_ids, default=0) + 1

    def _get_or_raise(self, template_id: int) -> Template:
        tmpl = self._templates.get_by_id(str(template_id))
        if tmpl is None:
            raise DomainValidationError("template_id", f"Template {template_id} not found.")
        return tmpl

    # ------------------------------------------------------------------ #
    # CRUD                                                                 #
    # ------------------------------------------------------------------ #

    def create_template(self, dto: CreateTemplateDto) -> TemplateDto:
        if not dto.name or not dto.name.strip():
            raise DomainValidationError("name", "Template name is required.")
        if not dto.html_content or not dto.html_content.strip():
            raise DomainValidationError("html_content", "HTML content is required.")
        if self._templates.get_by_name(dto.name.strip()) is not None:
            raise DomainValidationError("name", f"A template named '{dto.name}' already exists.")

        category = _category_from_str(dto.category)
        tmpl = Template.create(
            id=self._next_template_id(),
            name=dto.name.strip(),
            category=category,
            html_content=dto.html_content,
            created_by=dto.created_by,
            description=dto.description or None,
        )
        tmpl.is_published = dto.is_published
        for tag in dto.tags:
            tmpl.add_tag(tag)

        self._templates.add(tmpl)
        self._audit.log_event(
            event_type="TEMPLATE_CREATED",
            description=f"Template '{tmpl.name}' created",
            actor=dto.created_by,
            entity_id=str(tmpl.id),
            metadata={"category": category.value, "published": tmpl.is_published},
        )
        return _template_to_dto(tmpl)

    def get_template_by_id(self, template_id: int) -> Optional[TemplateDto]:
        tmpl = self._templates.get_by_id(str(template_id))
        return _template_to_dto(tmpl) if tmpl else None

    def get_template_by_name(self, name: str) -> Optional[TemplateDto]:
        tmpl = self._templates.get_by_name(name)
        return _template_to_dto(tmpl) if tmpl else None

    def get_all_templates(self) -> List[TemplateDto]:
        return [_template_to_dto(t) for t in self._templates.get_all()]

    def get_published_templates(self) -> List[TemplateDto]:
        return [_template_to_dto(t) for t in self._templates.get_published()]

    def update_template(self, dto: UpdateTemplateDto) -> TemplateDto:
        tmpl = self._get_or_raise(dto.template_id)
        if not dto.html_content or not dto.html_content.strip():
            raise DomainValidationError("html_content", "New HTML content cannot be empty.")

        tmpl.create_new_version(
            new_html=dto.html_content,
            changed_by=dto.updated_by,
            notes=dto.version_notes or None,
        )
        if dto.description is not None:
            tmpl.description = dto.description
        if dto.tags is not None:
            tmpl.tags = []
            for tag in dto.tags:
                tmpl.add_tag(tag)

        self._templates.update(tmpl)
        self._audit.log_event(
            event_type="TEMPLATE_UPDATED",
            description=f"Template '{tmpl.name}' updated to v{tmpl.version_number}",
            actor=dto.updated_by,
            entity_id=str(tmpl.id),
            metadata={"new_version": tmpl.version_number},
        )
        return _template_to_dto(tmpl)

    def delete_template(self, template_id: int, deleted_by: str) -> None:
        tmpl = self._get_or_raise(template_id)
        self._templates.delete(str(template_id))
        self._audit.log_event(
            event_type="TEMPLATE_DELETED",
            description=f"Template '{tmpl.name}' deleted",
            actor=deleted_by,
            entity_id=str(template_id),
            metadata={"name": tmpl.name},
        )

    def publish_template(self, template_id: int, published_by: str) -> None:
        tmpl = self._get_or_raise(template_id)
        tmpl.publish(published_by)
        self._templates.update(tmpl)
        self._audit.log_event(
            event_type="TEMPLATE_PUBLISHED",
            description=f"Template '{tmpl.name}' published",
            actor=published_by,
            entity_id=str(template_id),
        )

    def unpublish_template(self, template_id: int, unpublished_by: str) -> None:
        tmpl = self._get_or_raise(template_id)
        tmpl.unpublish(unpublished_by)
        self._templates.update(tmpl)
        self._audit.log_event(
            event_type="TEMPLATE_UNPUBLISHED",
            description=f"Template '{tmpl.name}' unpublished",
            actor=unpublished_by,
            entity_id=str(template_id),
        )

    # ------------------------------------------------------------------ #
    # Placeholders                                                         #
    # ------------------------------------------------------------------ #

    def add_placeholder(self, template_id: int, placeholder_dto: TemplatePlaceholderDto) -> None:
        tmpl = self._get_or_raise(template_id)
        from domain.entities.template_placeholder import PlaceholderType
        ph = TemplatePlaceholder(
            name=placeholder_dto.name,
            description=placeholder_dto.description,
            placeholder_type=PlaceholderType(placeholder_dto.placeholder_type),
            is_required=placeholder_dto.is_required,
            default_value=placeholder_dto.default_value,
            sample_value=placeholder_dto.sample_value,
        )
        tmpl.add_placeholder(ph)
        self._templates.update(tmpl)

    def remove_placeholder(self, template_id: int, placeholder_name: str) -> None:
        tmpl = self._get_or_raise(template_id)
        tmpl.remove_placeholder(placeholder_name)
        self._templates.update(tmpl)

    def get_placeholders(self, template_id: int) -> List[TemplatePlaceholderDto]:
        tmpl = self._get_or_raise(template_id)
        return _template_to_dto(tmpl).placeholders

    def auto_detect_placeholders(self, template_id: int) -> List[str]:
        tmpl = self._get_or_raise(template_id)
        new_names = tmpl.auto_detect_placeholders()
        if new_names:
            self._templates.update(tmpl)
        return new_names

    # ------------------------------------------------------------------ #
    # HTML validation                                                      #
    # ------------------------------------------------------------------ #

    def validate_template_html(self, html: str) -> List[str]:
        if not html or not html.strip():
            return ["HTML content is empty."]
        return _HTMLValidator().check(html)

    # ------------------------------------------------------------------ #
    # Rendering                                                            #
    # ------------------------------------------------------------------ #

    def render_template(self, template_id: int, values: Dict[str, Any]) -> str:
        tmpl = self._get_or_raise(template_id)
        if tmpl.placeholders:
            rendered = tmpl.render(values)
        else:
            # No registered placeholders — do simple {key} substitution
            rendered = tmpl.html_content
            for key, val in values.items():
                rendered = rendered.replace(f"{{{key}}}", str(val))
        tmpl.record_usage()
        self._templates.update(tmpl)
        return rendered

    def render_template_by_name(self, name: str, values: Dict[str, Any]) -> str:
        tmpl = self._templates.get_by_name(name)
        if tmpl is None:
            raise DomainValidationError("name", f"Template '{name}' not found.")
        if tmpl.placeholders:
            rendered = tmpl.render(values)
        else:
            rendered = tmpl.html_content
            for key, val in values.items():
                rendered = rendered.replace(f"{{{key}}}", str(val))
        tmpl.record_usage()
        self._templates.update(tmpl)
        return rendered

    def generate_preview(self, template_id: int) -> str:
        tmpl = self._get_or_raise(template_id)
        return tmpl.generate_preview()

    def validate_placeholder_values(self, template_id: int, values: Dict[str, Any]) -> List[str]:
        tmpl = self._get_or_raise(template_id)
        return tmpl.validate_placeholder_values(values)

    # ------------------------------------------------------------------ #
    # Category & filtering                                                 #
    # ------------------------------------------------------------------ #

    def get_templates_by_category(self, category: str) -> List[TemplateDto]:
        return [_template_to_dto(t) for t in self._templates.get_by_category(category)]

    def get_templates_by_tag(self, tag: str) -> List[TemplateDto]:
        tag_lower = tag.strip().lower()
        return [
            _template_to_dto(t)
            for t in self._templates.get_all()
            if tag_lower in [x.lower() for x in t.tags]
        ]

    def get_recent_templates(self, limit: int = 10) -> List[TemplateDto]:
        all_used = [t for t in self._templates.get_all() if t.last_used_at is not None]
        all_used.sort(key=lambda t: t.last_used_at, reverse=True)
        return [_template_to_dto(t) for t in all_used[:limit]]

    def get_most_used_templates(self, limit: int = 10) -> List[TemplateDto]:
        ranked = sorted(self._templates.get_all(), key=lambda t: t.usage_count, reverse=True)
        return [_template_to_dto(t) for t in ranked[:limit]]

    def get_favorite_templates(self, username: str) -> List[TemplateDto]:
        return [_template_to_dto(t) for t in self._templates.get_all() if t.is_favorite]

    def add_to_favorites(self, template_id: int, username: str) -> None:
        tmpl = self._get_or_raise(template_id)
        tmpl.is_favorite = True
        self._templates.update(tmpl)

    def remove_from_favorites(self, template_id: int, username: str) -> None:
        tmpl = self._get_or_raise(template_id)
        tmpl.is_favorite = False
        self._templates.update(tmpl)

    def search_templates(self, query: str) -> List[TemplateDto]:
        return [_template_to_dto(t) for t in self._templates.search(query)]

    # ------------------------------------------------------------------ #
    # Version history                                                      #
    # ------------------------------------------------------------------ #

    def get_version_history(self, template_id: int) -> List[TemplateVersionDto]:
        tmpl = self._get_or_raise(template_id)
        return [
            TemplateVersionDto(
                version_number=v.version_number,
                created_at=v.created_at,
                created_by=v.created_by,
                notes=v.notes or "",
            )
            for v in tmpl.version_history
        ]

    def get_template_version(self, template_id: int, version_number: int) -> Optional[TemplateVersionDto]:
        tmpl = self._get_or_raise(template_id)
        v = tmpl.get_version(version_number)
        if v is None:
            return None
        return TemplateVersionDto(
            version_number=v.version_number,
            created_at=v.created_at,
            created_by=v.created_by,
            notes=v.notes or "",
        )

    def rollback_to_version(self, template_id: int, version_number: int, rolled_back_by: str) -> None:
        tmpl = self._get_or_raise(template_id)
        tmpl.rollback_to_version(version_number, rolled_back_by)
        self._templates.update(tmpl)
        self._audit.log_event(
            event_type="TEMPLATE_ROLLED_BACK",
            description=f"Template '{tmpl.name}' rolled back to v{version_number}",
            actor=rolled_back_by,
            entity_id=str(template_id),
            metadata={"target_version": version_number, "new_version": tmpl.version_number},
        )

    # ------------------------------------------------------------------ #
    # Usage & statistics                                                   #
    # ------------------------------------------------------------------ #

    def record_usage(self, template_id: int) -> None:
        tmpl = self._get_or_raise(template_id)
        tmpl.record_usage()
        self._templates.update(tmpl)

    def get_statistics(self, template_id: int) -> dict:
        tmpl = self._get_or_raise(template_id)
        return {
            "template_id": tmpl.id,
            "name": tmpl.name,
            "usage_count": tmpl.usage_count,
            "last_used_at": tmpl.last_used_at.isoformat() if tmpl.last_used_at else None,
            "version_count": len(tmpl.version_history),
            "placeholder_count": tmpl.placeholder_count,
            "is_published": tmpl.is_published,
        }

    def get_all_statistics(self) -> List[dict]:
        return [self.get_statistics(t.id) for t in self._templates.get_all()]

    # ------------------------------------------------------------------ #
    # Import / export                                                      #
    # ------------------------------------------------------------------ #

    def import_from_docx(self, file_path: str, imported_by: str) -> TemplateDto:
        """Import a DOCX file as a new template (requires python-docx)."""
        try:
            import docx  # type: ignore[import-untyped]
        except ImportError:
            raise RuntimeError(
                "DOCX import requires the 'python-docx' package: pip install python-docx"
            )
        doc = docx.Document(file_path)
        html_parts = []
        for para in doc.paragraphs:
            if para.style.name.startswith("Heading"):
                level = para.style.name.split()[-1]
                html_parts.append(f"<h{level}>{para.text}</h{level}>")
            else:
                html_parts.append(f"<p>{para.text}</p>")
        html = "\n".join(html_parts)
        import os
        name = os.path.splitext(os.path.basename(file_path))[0]
        dto = CreateTemplateDto(
            name=name,
            category=TemplateCategory.BASIC.value,
            html_content=html,
            created_by=imported_by,
            description=f"Imported from {os.path.basename(file_path)}",
        )
        return self.create_template(dto)

    def export_to_docx(self, template_id: int, output_path: str) -> None:
        """Export template HTML to a DOCX file (requires python-docx)."""
        try:
            import docx  # type: ignore[import-untyped]
        except ImportError:
            raise RuntimeError(
                "DOCX export requires the 'python-docx' package: pip install python-docx"
            )
        tmpl = self._get_or_raise(template_id)
        doc = docx.Document()
        doc.add_heading(tmpl.name, 0)
        # Strip HTML tags for basic DOCX output
        plain = re.sub(r"<[^>]+>", " ", tmpl.html_content)
        plain = re.sub(r"\s+", " ", plain).strip()
        doc.add_paragraph(plain)
        doc.save(output_path)

    def import_from_html(self, html: str, name: str, imported_by: str) -> TemplateDto:
        dto = CreateTemplateDto(
            name=name,
            category=TemplateCategory.BASIC.value,
            html_content=html,
            created_by=imported_by,
        )
        return self.create_template(dto)

    def export_to_html(self, template_id: int) -> str:
        tmpl = self._get_or_raise(template_id)
        return tmpl.html_content

    def import_from_json(self, json_data: dict, imported_by: str) -> TemplateDto:
        dto = CreateTemplateDto(
            name=json_data["name"],
            category=json_data.get("category", TemplateCategory.BASIC.value),
            html_content=json_data["html_content"],
            created_by=imported_by,
            description=json_data.get("description", ""),
            tags=json_data.get("tags", []),
            is_published=bool(json_data.get("is_published", False)),
        )
        return self.create_template(dto)

    def export_to_json(self, template_id: int) -> dict:
        tmpl = self._get_or_raise(template_id)
        return {
            "name": tmpl.name,
            "category": tmpl.category.value,
            "description": tmpl.description or "",
            "html_content": tmpl.html_content,
            "version_number": tmpl.version_number,
            "is_published": tmpl.is_published,
            "tags": list(tmpl.tags),
            "placeholders": [
                {
                    "name": ph.name,
                    "type": ph.placeholder_type.value,
                    "description": ph.description,
                    "required": ph.is_required,
                }
                for ph in tmpl.placeholders
            ],
        }

    # ------------------------------------------------------------------ #
    # Seeding / backward-compat                                            #
    # ------------------------------------------------------------------ #

    def get_or_create_swgde_nist_template(self) -> TemplateDto:
        existing = self._templates.get_by_name("SWGDE/NIST Standard")
        if existing:
            return _template_to_dto(existing)
        from templates import DEFAULT_TEMPLATES
        html = DEFAULT_TEMPLATES.get("SWGDE/NIST Standard", "<h1>SWGDE/NIST Standard Report</h1>")
        dto = CreateTemplateDto(
            name="SWGDE/NIST Standard",
            category=TemplateCategory.SWGDE_NIST.value,
            html_content=html,
            created_by="system",
            description="Standard SWGDE/NIST digital forensic examination report",
            is_published=True,
        )
        return self.create_template(dto)

    def get_or_create_basic_template(self) -> TemplateDto:
        existing = self._templates.get_by_name("Basic Template")
        if existing:
            return _template_to_dto(existing)
        from templates import DEFAULT_TEMPLATES
        html = DEFAULT_TEMPLATES.get("Basic Template", "<h1>Forensic Report</h1>")
        dto = CreateTemplateDto(
            name="Basic Template",
            category=TemplateCategory.BASIC.value,
            html_content=html,
            created_by="system",
            description="Simple forensic report template",
            is_published=True,
        )
        return self.create_template(dto)

    def initialize_default_templates(self) -> None:
        """Seed default templates if none exist; import templates.json when present."""
        import json
        import os
        # If templates.json exists, import it (backward-compat)
        json_path = "templates.json"
        if os.path.exists(json_path):
            try:
                with open(json_path, encoding="utf-8") as fh:
                    saved: dict = json.load(fh)
                for name, html in saved.items():
                    if not html or not html.strip():
                        continue
                    if self._templates.get_by_name(name) is None:
                        cat = TemplateCategory.SWGDE_NIST if "swgde" in name.lower() or "nist" in name.lower() else TemplateCategory.BASIC
                        try:
                            self.create_template(
                                CreateTemplateDto(
                                    name=name,
                                    category=cat.value,
                                    html_content=html,
                                    created_by="system",
                                    is_published=True,
                                )
                            )
                        except DomainValidationError:
                            pass  # skip duplicates on re-run
            except Exception:
                pass  # graceful degradation

        # Always ensure the two built-in defaults exist
        self.get_or_create_swgde_nist_template()
        self.get_or_create_basic_template()

    def clone_template(self, template_id: int, new_name: str, cloned_by: str) -> TemplateDto:
        tmpl = self._get_or_raise(template_id)
        if not new_name or not new_name.strip():
            raise DomainValidationError("new_name", "Clone name cannot be empty.")
        dto = CreateTemplateDto(
            name=new_name.strip(),
            category=tmpl.category.value,
            html_content=tmpl.html_content,
            created_by=cloned_by,
            description=f"Clone of '{tmpl.name}'",
            tags=list(tmpl.tags),
            is_published=False,
        )
        cloned = self.create_template(dto)
        self._audit.log_event(
            event_type="TEMPLATE_CLONED",
            description=f"Template '{tmpl.name}' cloned as '{new_name}'",
            actor=cloned_by,
            entity_id=str(template_id),
            metadata={"new_template_id": cloned.template_id},
        )
        return cloned

    # ------------------------------------------------------------------ #
    # C# parity additions                                                  #
    # ------------------------------------------------------------------ #

    def get_or_create_mobile_forensics_template(self) -> TemplateDto:
        """Return (or create) the built-in Mobile Forensics template."""
        existing = [t for t in self._templates.get_all() if t.name == "Mobile Forensics Report"]
        if existing:
            return _template_to_dto(existing[0])
        dto = CreateTemplateDto(
            name="Mobile Forensics Report",
            category=TemplateCategory.MOBILE_DEVICE.value,
            html_content=(
                "<h1>Mobile Forensics Report</h1>"
                "<p><strong>Device Type:</strong> {{device_type}}</p>"
                "<p><strong>Device Identifier:</strong> {{device_id}}</p>"
                "<p><strong>Operating System:</strong> {{os_version}}</p>"
                "<p><strong>Acquisition Method:</strong> {{acquisition_method}}</p>"
                "<h2>Findings</h2><p>{{findings}}</p>"
                "<h2>Conclusions</h2><p>{{conclusions}}</p>"
            ),
            created_by="system",
            description="Standard Mobile Device Forensics Report template",
            tags=["mobile", "forensics", "built-in"],
            is_published=True,
        )
        return self.create_template(dto)

    def set_parent_template(
        self,
        template_id: int,
        parent_template_id: Optional[int],
        set_by: str,
    ) -> TemplateDto:
        """Set or clear the parent template for template inheritance."""
        tmpl = self._get_or_raise(template_id)
        if parent_template_id is not None:
            # Validate parent exists
            self._get_or_raise(parent_template_id)
            if parent_template_id == template_id:
                raise DomainValidationError(
                    "parent_template_id", "A template cannot be its own parent."
                )
        tmpl.parent_template_id = parent_template_id
        self._templates.update(tmpl)
        self._audit.log_event(
            event_type="TEMPLATE_PARENT_SET",
            description=(
                f"Template '{tmpl.name}' parent set to {parent_template_id!r}"
            ),
            actor=set_by,
            entity_id=str(template_id),
            metadata={"parent_template_id": parent_template_id},
        )
        return self._to_dto(tmpl)

    def export_templates_as_zip(self, template_ids: List[int], output_path: str) -> None:
        """Export a collection of templates to a ZIP archive."""
        import json as _json
        import zipfile
        import pathlib

        pathlib.Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        with zipfile.ZipFile(output_path, "w", zipfile.ZIP_DEFLATED) as zf:
            for tid in template_ids:
                try:
                    data = self.export_to_json(tid)
                    safe_name = re.sub(r"[^a-zA-Z0-9_-]", "_", data["name"])
                    zf.writestr(f"{safe_name}_{tid}.json", _json.dumps(data, indent=2))
                except Exception:
                    pass  # skip missing templates gracefully

    def import_templates_from_zip(self, zip_path: str, imported_by: str) -> List[TemplateDto]:
        """Import templates from a ZIP archive. Returns list of imported TemplateDto."""
        import json as _json
        import zipfile

        results: List[TemplateDto] = []
        with zipfile.ZipFile(zip_path, "r") as zf:
            for name in zf.namelist():
                if name.endswith(".json"):
                    try:
                        raw = zf.read(name).decode("utf-8")
                        data = _json.loads(raw)
                        tmpl = self.import_from_json(data, imported_by)
                        results.append(tmpl)
                    except Exception:
                        pass  # skip malformed entries gracefully
        return results

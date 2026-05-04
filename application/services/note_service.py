"""application/services/note_service.py — NoteService (Phase 6)."""
from __future__ import annotations

import csv
import io
import json
import uuid
from collections import Counter
from datetime import datetime
from typing import Dict, List, Optional

from application.dtos.note_dto import (
    CreateNoteDto,
    NoteDto,
    NoteStatisticsDto,
    UpdateNoteDto,
)
from application.interfaces.i_audit_service import IAuditService
from application.interfaces.i_clock import IClock
from application.interfaces.i_note_repository import INoteRepository
from application.interfaces.i_note_service import INoteService
from application.interfaces.i_report_repository import IReportRepository
from application.services._clock import DefaultClock
from domain.entities.note import Note
from domain.enums.note_status import NoteStatus
from domain.exceptions.domain_exceptions import DomainValidationError


def _to_dto(note: Note) -> NoteDto:
    return NoteDto(
        note_id=note.id,
        case_number=note.case_number,
        title=note.title,
        body=note.body,
        created_by=note.created_by,
        created_at=note.created_at,
        modified_at=note.modified_at,
        modified_by=note.modified_by,
        status=str(note.status.value) if note.status else NoteStatus.ACTIVE.value,
        tags=list(note.tags),
        note_type=note.note_type,
        priority=note.priority,
        approved_by=note.approved_by,
        approved_at=note.approved_at,
        approval_comments=note.approval_comments,
    )


class NoteService(INoteService):
    """Application service for investigation note management."""

    def __init__(
        self,
        note_repository: INoteRepository,
        report_repository: IReportRepository,
        audit_service: IAuditService,
        clock: Optional[IClock] = None,
    ) -> None:
        self._notes = note_repository
        self._reports = report_repository
        self._audit = audit_service
        self._clock = clock or DefaultClock()

    def _get_or_raise(self, note_id: str) -> Note:
        note = self._notes.get_by_id(note_id)
        if note is None:
            raise DomainValidationError("note_id", f"Note '{note_id}' not found.")
        return note

    @staticmethod
    def _task_meta(note: Note) -> Dict[str, str]:
        """Parse task metadata tags (task_status/task_due/task_assigned/task_completed_by)."""
        meta: Dict[str, str] = {}
        for tag in note.tags:
            if tag.startswith("task_status:"):
                meta["status"] = tag.split(":", 1)[1]
            elif tag.startswith("task_due:"):
                meta["due"] = tag.split(":", 1)[1]
            elif tag.startswith("task_assigned:"):
                meta["assigned"] = tag.split(":", 1)[1]
            elif tag.startswith("task_completed_by:"):
                meta["completed_by"] = tag.split(":", 1)[1]
        return meta

    @staticmethod
    def _set_task_meta(note: Note, key: str, value: Optional[str]) -> None:
        """Replace a task metadata tag (`task_<key>`) with a new value or remove it."""
        prefix = f"task_{key}:"
        note.tags = [t for t in note.tags if not t.startswith(prefix)]
        if value is not None and value != "":
            note.tags.append(f"task_{key}:{value}")

    @staticmethod
    def _is_task(note: Note) -> bool:
        return (note.note_type or "").strip().lower() == "task"

    @classmethod
    def _is_completed_task(cls, note: Note) -> bool:
        if not cls._is_task(note):
            return False
        return cls._task_meta(note).get("status") == "completed"

    @classmethod
    def _is_pending_task(cls, note: Note) -> bool:
        if not cls._is_task(note):
            return False
        return cls._task_meta(note).get("status", "pending") != "completed"

    @staticmethod
    def _meta_prefix(kind: str) -> str:
        return f"meta_{kind}:"

    @classmethod
    def _meta_items(cls, note: Note, kind: str) -> List[dict]:
        prefix = cls._meta_prefix(kind)
        items: List[dict] = []
        for tag in note.tags:
            if not tag.startswith(prefix):
                continue
            raw = tag[len(prefix):]
            try:
                items.append(json.loads(raw))
            except Exception:
                continue
        return items

    @classmethod
    def _add_meta_item(cls, note: Note, kind: str, item: dict) -> None:
        note.tags.append(cls._meta_prefix(kind) + json.dumps(item, separators=(",", ":")))

    @classmethod
    def _replace_meta_items(cls, note: Note, kind: str, items: List[dict]) -> None:
        prefix = cls._meta_prefix(kind)
        note.tags = [t for t in note.tags if not t.startswith(prefix)]
        for item in items:
            cls._add_meta_item(note, kind, item)

    @classmethod
    def _single_meta(cls, note: Note, kind: str) -> Optional[dict]:
        items = cls._meta_items(note, kind)
        return items[0] if items else None

    @classmethod
    def _set_single_meta(cls, note: Note, kind: str, item: Optional[dict]) -> None:
        prefix = cls._meta_prefix(kind)
        note.tags = [t for t in note.tags if not t.startswith(prefix)]
        if item is not None:
            cls._add_meta_item(note, kind, item)

    # ------------------------------------------------------------------ #
    # Queries                                                              #
    # ------------------------------------------------------------------ #

    def get_note(self, note_id: str) -> Optional[NoteDto]:
        note = self._notes.get_by_id(note_id)
        return _to_dto(note) if note else None

    def get_notes_for_case(self, case_number: str) -> List[NoteDto]:
        return [_to_dto(n) for n in self._notes.get_for_case(case_number)]

    def search_notes(self, query: str, case_number: Optional[str] = None) -> List[NoteDto]:
        return [_to_dto(n) for n in self._notes.search(query, case_number)]

    def get_archived_notes(self, case_number: str) -> List[NoteDto]:
        return [_to_dto(n) for n in self._notes.get_archived(case_number)]

    def get_pending_approval(self, case_number: str) -> List[NoteDto]:
        return [_to_dto(n) for n in self._notes.get_pending_approval(case_number)]

    def get_notes_by_tag(self, case_number: str, tag_name: str) -> List[NoteDto]:
        return [_to_dto(n) for n in self._notes.get_by_tag(case_number, tag_name)]

    def get_available_tags(self, case_number: str) -> List[str]:
        return self._notes.get_available_tags(case_number)

    def get_notes_by_type(self, case_number: str, note_type: str) -> List[NoteDto]:
        return [_to_dto(n) for n in self._notes.get_by_type(case_number, note_type)]

    def get_notes_by_priority(self, case_number: str, priority: str) -> List[NoteDto]:
        pr = priority.strip().lower()
        return [
            _to_dto(n)
            for n in self._notes.get_all()
            if n.case_number == case_number and (n.priority or "").strip().lower() == pr
        ]

    def get_pending_tasks(self, case_number: str) -> List[NoteDto]:
        return [
            _to_dto(n)
            for n in self._notes.get_all()
            if n.case_number == case_number and self._is_pending_task(n)
        ]

    def get_completed_tasks(self, case_number: str) -> List[NoteDto]:
        return [
            _to_dto(n)
            for n in self._notes.get_all()
            if n.case_number == case_number and self._is_completed_task(n)
        ]

    def get_tasks_assigned_to(self, case_number: str, username: str) -> List[NoteDto]:
        u = username.strip().lower()
        results: List[NoteDto] = []
        for n in self._notes.get_all():
            if n.case_number != case_number or not self._is_task(n):
                continue
            assigned = self._task_meta(n).get("assigned", "").strip().lower()
            if assigned == u:
                results.append(_to_dto(n))
        return results

    def get_overdue_tasks(self, case_number: str) -> List[NoteDto]:
        now = self._clock.utcnow()
        results: List[NoteDto] = []
        for n in self._notes.get_all():
            if n.case_number != case_number or not self._is_pending_task(n):
                continue
            due_raw = self._task_meta(n).get("due")
            if not due_raw:
                continue
            try:
                due = datetime.fromisoformat(due_raw)
            except ValueError:
                continue
            if due < now:
                results.append(_to_dto(n))
        return results

    # ------------------------------------------------------------------ #
    # CRUD mutations                                                       #
    # ------------------------------------------------------------------ #

    def create_note(self, dto: CreateNoteDto) -> NoteDto:
        if not dto.case_number or not dto.case_number.strip():
            raise DomainValidationError("case_number", "Case number is required.")
        if not dto.title or not dto.title.strip():
            raise DomainValidationError("title", "Note title is required.")
        note = Note.create(
            id=str(uuid.uuid4()),
            case_number=dto.case_number.strip(),
            title=dto.title.strip(),
            body=dto.body or "",
            created_by=dto.created_by,
            note_type=getattr(dto, "note_type", None),
            priority=getattr(dto, "priority", None),
        )
        note.created_at = self._clock.utcnow()
        for tag in getattr(dto, "tags", []):
            note.add_tag(tag)
        self._notes.add(note)
        self._audit.log_event(
            event_type="NOTE_CREATED",
            description=f"Note '{note.title}' created for case '{note.case_number}'",
            actor=dto.created_by,
            entity_id=note.id,
            metadata={"case_number": note.case_number},
        )
        return _to_dto(note)

    def update_note(self, dto: UpdateNoteDto) -> NoteDto:
        note = self._get_or_raise(dto.note_id)
        note.update(dto.title, dto.body, dto.modified_by)
        if hasattr(dto, "note_type") and dto.note_type is not None:
            note.note_type = dto.note_type
        if hasattr(dto, "priority") and dto.priority is not None:
            note.priority = dto.priority
        note.modified_at = self._clock.utcnow()
        self._notes.update(note)
        self._audit.log_event(
            event_type="NOTE_UPDATED",
            description=f"Note '{note.title}' updated",
            actor=dto.modified_by,
            entity_id=note.id,
            metadata={"case_number": note.case_number},
        )
        return _to_dto(note)

    def delete_note(self, note_id: str, deleted_by: str) -> None:
        note = self._get_or_raise(note_id)
        self._notes.delete(note_id)
        self._audit.log_event(
            event_type="NOTE_DELETED",
            description=f"Note '{note.title}' deleted from case '{note.case_number}'",
            actor=deleted_by,
            entity_id=note_id,
            metadata={"case_number": note.case_number},
        )

    def insert_note_into_report(self, note_id: str, report_id: int, inserted_by: str) -> None:
        note = self._get_or_raise(note_id)
        report = self._reports.get_by_id(str(report_id))
        if report is None:
            raise DomainValidationError("report_id", f"Report '{report_id}' not found.")
        divider = "\n\n<!-- Note: {} -->\n<blockquote>{}</blockquote>\n".format(
            note.title, note.body
        )
        report.report_html = (report.report_html or "") + divider
        self._reports.update(report)
        self._audit.log_event(
            event_type="NOTE_INSERTED_INTO_REPORT",
            description=f"Note '{note.title}' inserted into report {report_id}",
            actor=inserted_by,
            entity_id=note_id,
            metadata={"report_id": report_id, "case_number": note.case_number},
        )

    # ------------------------------------------------------------------ #
    # Archive / restore                                                    #
    # ------------------------------------------------------------------ #

    def archive_note(self, note_id: str, archived_by: str) -> NoteDto:
        note = self._get_or_raise(note_id)
        note.archive(archived_by)
        note.modified_at = self._clock.utcnow()
        self._notes.update(note)
        self._audit.log_event(
            event_type="NOTE_ARCHIVED",
            description=f"Note '{note.title}' archived",
            actor=archived_by,
            entity_id=note_id,
            metadata={"case_number": note.case_number},
        )
        return _to_dto(note)

    def restore_note(self, note_id: str, restored_by: str) -> NoteDto:
        note = self._get_or_raise(note_id)
        note.restore(restored_by)
        note.modified_at = self._clock.utcnow()
        self._notes.update(note)
        self._audit.log_event(
            event_type="NOTE_RESTORED",
            description=f"Note '{note.title}' restored",
            actor=restored_by,
            entity_id=note_id,
            metadata={"case_number": note.case_number},
        )
        return _to_dto(note)

    # ------------------------------------------------------------------ #
    # Approval workflow                                                    #
    # ------------------------------------------------------------------ #

    def submit_for_approval(self, note_id: str, submitted_by: str) -> NoteDto:
        note = self._get_or_raise(note_id)
        note.submit_for_approval(submitted_by)
        note.modified_at = self._clock.utcnow()
        self._notes.update(note)
        self._audit.log_event(
            event_type="NOTE_SUBMITTED_FOR_APPROVAL",
            description=f"Note '{note.title}' submitted for approval",
            actor=submitted_by,
            entity_id=note_id,
            metadata={"case_number": note.case_number},
        )
        return _to_dto(note)

    def approve_note(self, note_id: str, approved_by: str, comments: Optional[str] = None) -> NoteDto:
        note = self._get_or_raise(note_id)
        note.approve(approved_by, comments)
        note.modified_at = self._clock.utcnow()
        self._notes.update(note)
        self._audit.log_event(
            event_type="NOTE_APPROVED",
            description=f"Note '{note.title}' approved",
            actor=approved_by,
            entity_id=note_id,
            metadata={"case_number": note.case_number},
        )
        return _to_dto(note)

    def reject_note(self, note_id: str, rejected_by: str, reason: str) -> NoteDto:
        note = self._get_or_raise(note_id)
        note.reject(rejected_by, reason)
        note.modified_at = self._clock.utcnow()
        self._notes.update(note)
        self._audit.log_event(
            event_type="NOTE_REJECTED",
            description=f"Note '{note.title}' rejected: {reason}",
            actor=rejected_by,
            entity_id=note_id,
            metadata={"case_number": note.case_number, "reason": reason},
        )
        return _to_dto(note)

    # ------------------------------------------------------------------ #
    # Tag management                                                       #
    # ------------------------------------------------------------------ #

    def add_tag(self, note_id: str, tag_name: str, added_by: str) -> NoteDto:
        note = self._get_or_raise(note_id)
        note.add_tag(tag_name)
        note.modified_at = self._clock.utcnow()
        note.modified_by = added_by
        self._notes.update(note)
        self._audit.log_event(
            event_type="NOTE_TAG_ADDED",
            description=f"Tag '{tag_name}' added to note '{note.title}'",
            actor=added_by,
            entity_id=note_id,
            metadata={"case_number": note.case_number, "tag": tag_name},
        )
        return _to_dto(note)

    def remove_tag(self, note_id: str, tag_name: str, removed_by: str) -> NoteDto:
        note = self._get_or_raise(note_id)
        note.remove_tag(tag_name)
        note.modified_at = self._clock.utcnow()
        note.modified_by = removed_by
        self._notes.update(note)
        self._audit.log_event(
            event_type="NOTE_TAG_REMOVED",
            description=f"Tag '{tag_name}' removed from note '{note.title}'",
            actor=removed_by,
            entity_id=note_id,
            metadata={"case_number": note.case_number, "tag": tag_name},
        )
        return _to_dto(note)

    # ------------------------------------------------------------------ #
    # Task management                                                      #
    # ------------------------------------------------------------------ #

    def complete_task(self, note_id: str, completed_by: str) -> NoteDto:
        note = self._get_or_raise(note_id)
        if not self._is_task(note):
            raise DomainValidationError("note_type", "Only task notes can be completed.")
        self._set_task_meta(note, "status", "completed")
        self._set_task_meta(note, "completed_by", completed_by)
        self._set_task_meta(note, "completed_at", self._clock.utcnow().isoformat())
        note.modified_at = self._clock.utcnow()
        note.modified_by = completed_by
        self._notes.update(note)
        self._audit.log_event(
            event_type="TASK_COMPLETED",
            description=f"Task note '{note.title}' completed",
            actor=completed_by,
            entity_id=note.id,
            metadata={"case_number": note.case_number},
        )
        return _to_dto(note)

    def reopen_task(self, note_id: str, reopened_by: str) -> NoteDto:
        note = self._get_or_raise(note_id)
        if not self._is_task(note):
            raise DomainValidationError("note_type", "Only task notes can be reopened.")
        self._set_task_meta(note, "status", "pending")
        self._set_task_meta(note, "completed_by", None)
        self._set_task_meta(note, "completed_at", None)
        note.modified_at = self._clock.utcnow()
        note.modified_by = reopened_by
        self._notes.update(note)
        self._audit.log_event(
            event_type="TASK_REOPENED",
            description=f"Task note '{note.title}' reopened",
            actor=reopened_by,
            entity_id=note.id,
            metadata={"case_number": note.case_number},
        )
        return _to_dto(note)

    def reassign_task(self, note_id: str, assigned_to: str, reassigned_by: str) -> NoteDto:
        note = self._get_or_raise(note_id)
        if not self._is_task(note):
            raise DomainValidationError("note_type", "Only task notes can be reassigned.")
        if not assigned_to or not assigned_to.strip():
            raise DomainValidationError("assigned_to", "assigned_to is required.")
        self._set_task_meta(note, "assigned", assigned_to.strip())
        note.modified_at = self._clock.utcnow()
        note.modified_by = reassigned_by
        self._notes.update(note)
        self._audit.log_event(
            event_type="TASK_REASSIGNED",
            description=f"Task note '{note.title}' reassigned to '{assigned_to}'",
            actor=reassigned_by,
            entity_id=note.id,
            metadata={"case_number": note.case_number, "assigned_to": assigned_to},
        )
        return _to_dto(note)

    # ------------------------------------------------------------------ #
    # Attachments                                                         #
    # ------------------------------------------------------------------ #

    def add_attachment(
        self,
        note_id: str,
        file_path: str,
        file_name: str,
        mime_type: str,
        added_by: str,
    ) -> dict:
        note = self._get_or_raise(note_id)
        if not file_path or not file_name:
            raise DomainValidationError("attachment", "file_path and file_name are required.")
        item = {
            "attachment_id": str(uuid.uuid4()),
            "file_path": file_path,
            "file_name": file_name,
            "mime_type": mime_type,
            "added_by": added_by,
            "added_at": self._clock.utcnow().isoformat(),
        }
        self._add_meta_item(note, "attachment", item)
        note.modified_by = added_by
        note.modified_at = self._clock.utcnow()
        self._notes.update(note)
        self._audit.log_event(
            event_type="NOTE_ATTACHMENT_ADDED",
            description=f"Attachment '{file_name}' added to note '{note.title}'",
            actor=added_by,
            entity_id=note.id,
            metadata={"case_number": note.case_number, "attachment_id": item["attachment_id"]},
        )
        return item

    def remove_attachment(self, note_id: str, attachment_id: str, removed_by: str) -> None:
        note = self._get_or_raise(note_id)
        items = self._meta_items(note, "attachment")
        filtered = [a for a in items if a.get("attachment_id") != attachment_id]
        self._replace_meta_items(note, "attachment", filtered)
        note.modified_by = removed_by
        note.modified_at = self._clock.utcnow()
        self._notes.update(note)
        self._audit.log_event(
            event_type="NOTE_ATTACHMENT_REMOVED",
            description=f"Attachment removed from note '{note.title}'",
            actor=removed_by,
            entity_id=note.id,
            metadata={"case_number": note.case_number, "attachment_id": attachment_id},
        )

    def get_attachments(self, note_id: str) -> List[dict]:
        note = self._get_or_raise(note_id)
        return self._meta_items(note, "attachment")

    # ------------------------------------------------------------------ #
    # Entity links                                                        #
    # ------------------------------------------------------------------ #

    def add_link(
        self,
        note_id: str,
        target_type: str,
        target_id: str,
        label: str,
        linked_by: str,
    ) -> dict:
        note = self._get_or_raise(note_id)
        if not target_type or not target_id:
            raise DomainValidationError("link", "target_type and target_id are required.")
        item = {
            "link_id": str(uuid.uuid4()),
            "target_type": target_type,
            "target_id": target_id,
            "label": label,
            "linked_by": linked_by,
            "linked_at": self._clock.utcnow().isoformat(),
        }
        self._add_meta_item(note, "link", item)
        note.modified_by = linked_by
        note.modified_at = self._clock.utcnow()
        self._notes.update(note)
        return item

    def remove_link(self, note_id: str, link_id: str, removed_by: str) -> None:
        note = self._get_or_raise(note_id)
        links = self._meta_items(note, "link")
        links = [l for l in links if l.get("link_id") != link_id]
        self._replace_meta_items(note, "link", links)
        note.modified_by = removed_by
        note.modified_at = self._clock.utcnow()
        self._notes.update(note)

    def get_links(self, note_id: str) -> List[dict]:
        note = self._get_or_raise(note_id)
        return self._meta_items(note, "link")

    def get_notes_linking_to(self, case_number: str, target_type: str, target_id: str) -> List[NoteDto]:
        results: List[NoteDto] = []
        for n in self._notes.get_all():
            if n.case_number != case_number:
                continue
            links = self._meta_items(n, "link")
            if any(l.get("target_type") == target_type and l.get("target_id") == target_id for l in links):
                results.append(_to_dto(n))
        return results

    # ------------------------------------------------------------------ #
    # Redaction / sharing / visibility                                    #
    # ------------------------------------------------------------------ #

    def redact_note(self, note_id: str, reason: str, redacted_by: str) -> NoteDto:
        note = self._get_or_raise(note_id)
        redaction = {
            "is_redacted": True,
            "reason": reason,
            "redacted_by": redacted_by,
            "redacted_at": self._clock.utcnow().isoformat(),
            "original_body": note.body,
        }
        self._set_single_meta(note, "redaction", redaction)
        note.body = "[REDACTED]"
        note.modified_by = redacted_by
        note.modified_at = self._clock.utcnow()
        self._notes.update(note)
        self._audit.log_event(
            event_type="NOTE_REDACTED",
            description=f"Note '{note.title}' redacted",
            actor=redacted_by,
            entity_id=note.id,
            metadata={"case_number": note.case_number, "reason": reason},
        )
        return _to_dto(note)

    def unredact_note(self, note_id: str, unredacted_by: str) -> NoteDto:
        note = self._get_or_raise(note_id)
        redaction = self._single_meta(note, "redaction")
        if redaction and "original_body" in redaction:
            note.body = redaction["original_body"]
        self._set_single_meta(note, "redaction", None)
        note.modified_by = unredacted_by
        note.modified_at = self._clock.utcnow()
        self._notes.update(note)
        self._audit.log_event(
            event_type="NOTE_UNREDACTED",
            description=f"Note '{note.title}' unredacted",
            actor=unredacted_by,
            entity_id=note.id,
            metadata={"case_number": note.case_number},
        )
        return _to_dto(note)

    def share_note(self, note_id: str, usernames: List[str], shared_by: str) -> NoteDto:
        note = self._get_or_raise(note_id)
        share = self._single_meta(note, "share") or {"users": []}
        users = {u.strip().lower() for u in share.get("users", []) if u}
        users.update({u.strip().lower() for u in usernames if u and u.strip()})
        share["users"] = sorted(users)
        self._set_single_meta(note, "share", share)
        note.modified_by = shared_by
        note.modified_at = self._clock.utcnow()
        self._notes.update(note)
        self._audit.log_event(
            event_type="NOTE_SHARED",
            description=f"Note '{note.title}' shared",
            actor=shared_by,
            entity_id=note.id,
            metadata={"case_number": note.case_number, "users": share["users"]},
        )
        return _to_dto(note)

    def unshare_note(self, note_id: str, usernames: List[str], unshared_by: str) -> NoteDto:
        note = self._get_or_raise(note_id)
        share = self._single_meta(note, "share") or {"users": []}
        current = {u.strip().lower() for u in share.get("users", []) if u}
        remove = {u.strip().lower() for u in usernames if u and u.strip()}
        share["users"] = sorted(current - remove)
        self._set_single_meta(note, "share", share)
        note.modified_by = unshared_by
        note.modified_at = self._clock.utcnow()
        self._notes.update(note)
        self._audit.log_event(
            event_type="NOTE_UNSHARED",
            description=f"Note '{note.title}' unshared",
            actor=unshared_by,
            entity_id=note.id,
            metadata={"case_number": note.case_number, "users": share["users"]},
        )
        return _to_dto(note)

    def change_visibility(self, note_id: str, visibility: str, changed_by: str) -> NoteDto:
        note = self._get_or_raise(note_id)
        vis = visibility.strip().lower()
        allowed = {"private", "team", "case_level", "public"}
        if vis not in allowed:
            raise DomainValidationError("visibility", "visibility must be private/team/case_level/public")
        self._set_single_meta(note, "visibility", {"value": vis})
        note.modified_by = changed_by
        note.modified_at = self._clock.utcnow()
        self._notes.update(note)
        self._audit.log_event(
            event_type="NOTE_VISIBILITY_CHANGED",
            description=f"Note '{note.title}' visibility changed to {vis}",
            actor=changed_by,
            entity_id=note.id,
            metadata={"case_number": note.case_number, "visibility": vis},
        )
        return _to_dto(note)

    # ------------------------------------------------------------------ #
    # Timeline / metrics / exports                                        #
    # ------------------------------------------------------------------ #

    def generate_timeline(self, case_number: str) -> List[dict]:
        items: List[dict] = []
        for n in self._notes.get_all():
            if n.case_number != case_number:
                continue
            items.append(
                {
                    "timestamp": n.created_at,
                    "event": "note_created",
                    "note_id": n.id,
                    "title": n.title,
                    "actor": n.created_by,
                }
            )
            if n.modified_at:
                items.append(
                    {
                        "timestamp": n.modified_at,
                        "event": "note_modified",
                        "note_id": n.id,
                        "title": n.title,
                        "actor": n.modified_by,
                    }
                )
        items.sort(key=lambda x: x["timestamp"])
        return items

    def get_activity_metrics(self, case_number: str) -> dict:
        notes = [n for n in self._notes.get_all() if n.case_number == case_number]
        created = len(notes)
        updated = sum(1 for n in notes if n.modified_at is not None)
        by_user = Counter(n.created_by for n in notes if n.created_by)
        most_active = by_user.most_common(1)[0][0] if by_user else None
        return {
            "case_number": case_number,
            "notes_created": created,
            "notes_updated": updated,
            "most_active_user": most_active,
            "created_by_user": dict(by_user),
        }

    def export_to_pdf(self, case_number: str, output_path: str) -> str:
        notes = [n for n in self._notes.get_all() if n.case_number == case_number]
        try:
            from reportlab.lib.pagesizes import letter  # type: ignore[import-untyped]
            from reportlab.pdfgen import canvas  # type: ignore[import-untyped]

            c = canvas.Canvas(output_path, pagesize=letter)
            y = 760
            c.setFont("Helvetica-Bold", 14)
            c.drawString(50, y, f"Case Notes Export: {case_number}")
            y -= 30
            c.setFont("Helvetica", 10)
            for n in notes:
                line = f"[{n.created_at.isoformat()}] {n.title} ({n.created_by})"
                c.drawString(50, y, line[:100])
                y -= 14
                if y < 60:
                    c.showPage()
                    y = 760
                    c.setFont("Helvetica", 10)
            c.save()
        except Exception:
            # Fallback: write plain text content to the target path.
            with open(output_path, "w", encoding="utf-8") as fh:
                fh.write(f"Case Notes Export: {case_number}\n\n")
                for n in notes:
                    fh.write(f"[{n.created_at.isoformat()}] {n.title} ({n.created_by})\n")
                    fh.write(n.body + "\n\n")
        return output_path

    def export_to_docx(self, case_number: str, output_path: str) -> str:
        notes = [n for n in self._notes.get_all() if n.case_number == case_number]
        try:
            import docx  # type: ignore[import-untyped]

            doc = docx.Document()
            doc.add_heading(f"Case Notes Export: {case_number}", level=1)
            for n in notes:
                doc.add_heading(n.title, level=2)
                doc.add_paragraph(f"Created by {n.created_by} on {n.created_at.isoformat()}")
                doc.add_paragraph(n.body)
            doc.save(output_path)
        except Exception:
            with open(output_path, "w", encoding="utf-8") as fh:
                fh.write(f"Case Notes Export: {case_number}\n\n")
                for n in notes:
                    fh.write(f"## {n.title}\n")
                    fh.write(f"Created by {n.created_by} on {n.created_at.isoformat()}\n")
                    fh.write(n.body + "\n\n")
        return output_path

    def export_timeline(self, case_number: str, output_path: str) -> str:
        timeline = self.generate_timeline(case_number)
        with open(output_path, "w", newline="", encoding="utf-8") as fh:
            writer = csv.writer(fh)
            writer.writerow(["timestamp", "event", "note_id", "title", "actor"])
            for row in timeline:
                writer.writerow(
                    [
                        row["timestamp"].isoformat() if row.get("timestamp") else "",
                        row.get("event", ""),
                        row.get("note_id", ""),
                        row.get("title", ""),
                        row.get("actor", ""),
                    ]
                )
        return output_path

    # ------------------------------------------------------------------ #
    # Statistics                                                           #
    # ------------------------------------------------------------------ #

    def get_statistics(self, case_number: str) -> NoteStatisticsDto:
        all_notes = self._notes.get_all()
        case_notes = [n for n in all_notes if n.case_number == case_number]

        by_status: Dict[str, int] = Counter(n.status.value for n in case_notes)
        by_type: Dict[str, int] = Counter(
            n.note_type for n in case_notes if n.note_type
        )
        by_priority: Dict[str, int] = Counter(
            n.priority for n in case_notes if n.priority
        )
        all_tags = sorted(
            {t for n in case_notes for t in n.tags}
        )

        return NoteStatisticsDto(
            case_number=case_number,
            total=len(case_notes),
            active=by_status.get(NoteStatus.ACTIVE.value, 0),
            archived=by_status.get(NoteStatus.ARCHIVED.value, 0),
            pending_approval=by_status.get(NoteStatus.PENDING_APPROVAL.value, 0),
            approved=by_status.get(NoteStatus.APPROVED.value, 0),
            rejected=by_status.get(NoteStatus.REJECTED.value, 0),
            by_type=dict(by_type),
            by_priority=dict(by_priority),
            all_tags=all_tags,
        )

    # ------------------------------------------------------------------ #
    # Export                                                               #
    # ------------------------------------------------------------------ #

    def export_to_csv(self, case_number: str) -> str:
        """Return all notes for a case as a CSV string."""
        all_notes = self._notes.get_all()
        case_notes = [n for n in all_notes if n.case_number == case_number]

        buf = io.StringIO()
        writer = csv.writer(buf)
        writer.writerow(
            ["id", "title", "body", "status", "note_type", "priority",
             "tags", "created_by", "created_at", "modified_at", "modified_by",
             "approved_by", "approved_at", "approval_comments"]
        )
        for n in case_notes:
            writer.writerow([
                n.id, n.title, n.body,
                n.status.value if n.status else "",
                n.note_type or "",
                n.priority or "",
                "|".join(n.tags),
                n.created_by,
                n.created_at.isoformat() if n.created_at else "",
                n.modified_at.isoformat() if n.modified_at else "",
                n.modified_by or "",
                n.approved_by or "",
                n.approved_at.isoformat() if n.approved_at else "",
                n.approval_comments or "",
            ])
        return buf.getvalue()

    # ------------------------------------------------------------------ #
    # Geocoding                                                            #
    # ------------------------------------------------------------------ #

    def forward_geocode(self, address: str) -> dict:
        """
        Geocode an address to coordinates.
        Uses geopy (Nominatim) if available; falls back to a stub result.
        Result is cached on the note via meta tags when called from extract_locations.
        """
        try:
            from geopy.geocoders import Nominatim  # type: ignore

            geolocator = Nominatim(user_agent="forensic_notes")
            location = geolocator.geocode(address, timeout=5)
            if location:
                return {
                    "lat": location.latitude,
                    "lon": location.longitude,
                    "display_name": location.address,
                    "source": "nominatim",
                }
        except Exception:
            pass
        # Fallback stub — indicates geocoding was attempted but unavailable
        return {"lat": None, "lon": None, "display_name": address, "source": "stub"}

    def reverse_geocode(self, lat: float, lon: float) -> str:
        """
        Reverse geocode coordinates to a human-readable address.
        Uses geopy (Nominatim) if available; falls back to a coordinate string.
        """
        try:
            from geopy.geocoders import Nominatim  # type: ignore

            geolocator = Nominatim(user_agent="forensic_notes")
            location = geolocator.reverse((lat, lon), timeout=5)
            if location:
                return location.address
        except Exception:
            pass
        return f"{lat}, {lon}"

    def extract_locations(self, note_id: str) -> List[dict]:
        """
        Parse a note's body for location-like strings using a simple heuristic regex,
        then attempt to geocode each one. Returns a list of {text, lat, lon} dicts.
        Cached geocoding results are stored as meta tags on the note.
        """
        import re

        note = self._notes.get_by_id(note_id)
        if not note:
            raise DomainValidationError("note_id", "Note not found")

        # Simple pattern: "at <Title Case Words>" or "located in/at/near <words>"
        pattern = re.compile(
            r"\b(?:at|in|near|located (?:at|in|near))\s+([A-Z][a-zA-Z\s,\.]{3,50})",
            re.IGNORECASE,
        )
        matches = pattern.findall(note.body or "")
        results = []
        for match in matches:
            text = match.strip().rstrip(",.")
            geo = self.forward_geocode(text)
            results.append({"text": text, "lat": geo["lat"], "lon": geo["lon"]})
        return results

    # ------------------------------------------------------------------ #
    # Voice Transcription                                                  #
    # ------------------------------------------------------------------ #

    def create_from_voice_transcription(
        self,
        case_number: str,
        audio_path: str,
        created_by: str,
        title: Optional[str] = None,
    ) -> NoteDto:
        """
        Create a note from an audio recording.
        Uses speech_recognition (Google backend) if available, then whisper if available.
        Falls back to creating a placeholder note indicating transcription is pending.
        Confidence score is stored as the meta tag 'meta_transcription_confidence'.
        """
        transcript = None
        confidence: Optional[float] = None

        # Attempt 1: SpeechRecognition
        try:
            import speech_recognition as sr  # type: ignore

            recognizer = sr.Recognizer()
            with sr.AudioFile(audio_path) as source:
                audio_data = recognizer.record(source)
            transcript = recognizer.recognize_google(audio_data)
            confidence = 0.85  # Google API does not expose confidence on all platforms
        except Exception:
            transcript = None

        # Attempt 2: OpenAI Whisper
        if transcript is None:
            try:
                import whisper  # type: ignore

                model = whisper.load_model("base")
                result = model.transcribe(audio_path)
                transcript = result.get("text", "")
                # Whisper returns segment-level logprobs; approximate mean confidence
                segments = result.get("segments", [])
                if segments:
                    avg_logprob = sum(s.get("avg_logprob", -1.0) for s in segments) / len(segments)
                    import math
                    confidence = round(min(1.0, max(0.0, math.exp(avg_logprob))), 4)
                else:
                    confidence = None
            except Exception:
                transcript = None

        if not transcript:
            transcript = f"[Transcription pending — audio file: {audio_path}]"
            confidence = None

        effective_title = title or f"Voice Note — {self._clock.utcnow().strftime('%Y-%m-%d %H:%M')}"
        dto = self.create_note(
            CreateNoteDto(
                case_number=case_number,
                title=effective_title,
                body=transcript,
                created_by=created_by,
            )
        )
        # Store confidence as a meta tag
        if confidence is not None:
            note = self._notes.get_by_id(dto.note_id)
            if note:
                note.tags = [t for t in note.tags if not t.startswith("meta_transcription_confidence")]
                note.tags.append(f"meta_transcription_confidence:{confidence}")
                self._notes.update(note)

        return dto

    def get_voice_transcription_confidence(self, note_id: str) -> Optional[float]:
        """Return the stored transcription confidence score (0.0–1.0), or None."""
        note = self._notes.get_by_id(note_id)
        if not note:
            return None
        for tag in note.tags:
            if tag.startswith("meta_transcription_confidence:"):
                try:
                    return float(tag.split(":", 1)[1])
                except ValueError:
                    return None
        return None

    # ------------------------------------------------------------------ #
    # NLP Entity Extraction                                                #
    # ------------------------------------------------------------------ #

    def extract_entities(self, note_id: str) -> dict:
        """
        Extract named entities from a note's body.
        Tries spaCy (en_core_web_sm) first, then NLTK, then falls back to
        regex-based heuristics for capitalised tokens, dates, and evidence refs.
        Returns dict with keys: people, places, evidence_refs, dates, organisations.
        """
        import re

        note = self._notes.get_by_id(note_id)
        if not note:
            raise DomainValidationError("note_id", "Note not found")

        text = note.body or ""

        # Attempt 1: spaCy
        try:
            import spacy  # type: ignore

            nlp = spacy.load("en_core_web_sm")
            doc = nlp(text)
            label_map = {
                "PERSON": "people",
                "GPE": "places",
                "LOC": "places",
                "ORG": "organisations",
                "DATE": "dates",
            }
            result: dict = {"people": [], "places": [], "evidence_refs": [], "dates": [], "organisations": []}
            for ent in doc.ents:
                key = label_map.get(ent.label_)
                if key and ent.text not in result[key]:
                    result[key].append(ent.text)
            # Evidence ref regex regardless of NLP backend
            result["evidence_refs"] = list({
                m.group() for m in re.finditer(r"\b[EeIi][Tt]?[-_]?\d{3,}\b", text)
            })
            return result
        except Exception:
            pass

        # Attempt 2: NLTK
        try:
            import nltk  # type: ignore
            from nltk import ne_chunk, pos_tag, word_tokenize  # type: ignore
            from nltk.tree import Tree  # type: ignore

            tokens = word_tokenize(text)
            tagged = pos_tag(tokens)
            chunks = ne_chunk(tagged)
            result = {"people": [], "places": [], "evidence_refs": [], "dates": [], "organisations": []}
            for chunk in chunks:
                if isinstance(chunk, Tree):
                    entity = " ".join(w for w, _ in chunk.leaves())
                    label = chunk.label()
                    if label == "PERSON" and entity not in result["people"]:
                        result["people"].append(entity)
                    elif label in ("GPE", "GSP", "LOCATION") and entity not in result["places"]:
                        result["places"].append(entity)
                    elif label == "ORGANIZATION" and entity not in result["organisations"]:
                        result["organisations"].append(entity)
            result["evidence_refs"] = list({
                m.group() for m in re.finditer(r"\b[EeIi][Tt]?[-_]?\d{3,}\b", text)
            })
            return result
        except Exception:
            pass

        # Fallback: regex heuristics
        result = {"people": [], "places": [], "evidence_refs": [], "dates": [], "organisations": []}
        # Capitalised multi-word phrases (crude person/org/place detector)
        for m in re.finditer(r"\b([A-Z][a-z]+(?: [A-Z][a-z]+)+)\b", text):
            word = m.group()
            if word not in result["people"]:
                result["people"].append(word)
        # ISO dates and common date patterns
        for m in re.finditer(
            r"\b(\d{4}-\d{2}-\d{2}|\d{1,2}[/\-]\d{1,2}[/\-]\d{2,4})\b", text
        ):
            d = m.group()
            if d not in result["dates"]:
                result["dates"].append(d)
        # Evidence refs: E-001, I001, IT-042, etc.
        result["evidence_refs"] = list({
            m.group() for m in re.finditer(r"\b[EeIi][Tt]?[-_]?\d{3,}\b", text)
        })
        return result



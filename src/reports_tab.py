# reports_tab.py
import os
import shutil
import json
import hashlib
import logging
from datetime import datetime

# Initialize logger
logger = logging.getLogger(__name__)

try:
    from zoneinfo import ZoneInfo
except ImportError as e:
    logger.warning(f"Failed to import zoneinfo: {e}")
    ZoneInfo = None
try:
    from dateutil import parser as dateutil_parser
    from dateutil import tz as dateutil_tz
except Exception as e:
    logger.warning(f"Failed to import dateutil: {e}")
    dateutil_parser = None
    dateutil_tz = None
from PyQt5.QtWidgets import (
    QWidget, QVBoxLayout, QHBoxLayout, QLabel, QPushButton, QTextEdit, QToolBar, QAction,
    QFileDialog, QMessageBox, QInputDialog, QCheckBox, QSplitter, QTableWidget, QTableWidgetItem,
    QLineEdit, QListWidget, QListWidgetItem, QMainWindow, QComboBox, QColorDialog, QFontComboBox, QSpinBox,
    QDialog, QFormLayout, QDialogButtonBox, QFrame, QMenu, QToolButton, QSizePolicy
)
from PyQt5.QtGui import QTextCharFormat, QTextListFormat, QFont, QKeySequence, QPixmap
from PyQt5.QtCore import Qt

from security import compute_sha256
from templates import TemplateManager
from glossary_assist import GlossaryAssist
from peer_review import PeerReview
from base_editor import BaseEditor
from word_processor import WordProcessor

logger = logging.getLogger(__name__)


class ReportsTab(BaseEditor):
    def __init__(self, case_data, db_manager, audit_logger, current_user, parent=None):
        super().__init__(case_data, db_manager, audit_logger, parent)
        self.current_user = current_user
        self.case_dir = os.path.join("cases", case_data["case_number"])
        os.makedirs(self.case_dir, exist_ok=True)

        # Data structures
        self.appendices = []
        self.saved_pdf_hash = ""

        # Initialize advanced word processor
        self.word_processor = WordProcessor(self.editor)

        self.load_report()

    def setup_ui(self):
        """Override parent's setup_ui with Office-style layout"""
        # Initialize parent class attributes
        self.toolbar = None
        self.splitter = None
        self.editor_frame = None

        # Create main layout
        layout = QVBoxLayout(self)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.setSpacing(0)

        # ── Office-style formatting toolbar ──────────────────────────────
        # create_document_canvas() sets self.editor; do it first so toolbar
        # signals find a valid self.editor when they fire.
        canvas = self.create_document_canvas()

        # Alias so existing methods that reference self.report_editor keep working
        self.report_editor = self.editor

        # Wire up report-specific signals
        self.report_editor.setContextMenuPolicy(Qt.CustomContextMenu)
        self.report_editor.customContextMenuRequested.connect(self.on_report_context_menu)
        self.report_editor.textChanged.connect(self.on_text_changed)

        self.toolbar = self.create_toolbar()
        layout.addWidget(self.toolbar)
        layout.addWidget(canvas, stretch=1)

        # ── Footer: PDF hash label ────────────────────────────────────────
        self.hash_label = QLabel("Final PDF Hash: Not generated")
        self.hash_label.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Fixed)
        layout.addWidget(self.hash_label)

        # ── Status bar ────────────────────────────────────────────────────
        from base_editor import StatusBar
        self.status_bar = StatusBar()
        self.status_bar.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Fixed)
        layout.addWidget(self.status_bar)

    def load_report(self):
        html, appendices, pdf_hash = self.db.load_report(self.case_data['case_number'])
        if html:
            self.report_editor.setHtml(html)
        self.appendices = appendices
        if pdf_hash:
            self.hash_label.setText(f"Final PDF Hash (SHA-256):\n{pdf_hash}")
            self.saved_pdf_hash = pdf_hash
        self.update_appendices_list()

    def on_report_context_menu(self, pos):
        try:
            from auth import load_config
            cfg = load_config()
        except Exception as e:
            logger.warning(f"Failed to load config for context menu: {e}")
            cfg = {}

        rpt_cfg = cfg.get('context_menu', {}).get('reports', {}) if cfg else {}

        menu = self.report_editor.createStandardContextMenu()
        menu.addSeparator()

        if rpt_cfg.get('insert_timestamp', True):
            insert_local = QAction('Insert Date/Time (Local)', self)
            insert_local.setShortcut(QKeySequence('Ctrl+Shift+T'))
            insert_local.triggered.connect(lambda: self.insert_timestamp_at_cursor(fmt='local'))
            insert_iso = QAction('Insert ISO Date/Time (with TZ)', self)
            insert_iso.setShortcut(QKeySequence('Ctrl+Alt+Shift+T'))
            insert_iso.triggered.connect(lambda: self.insert_timestamp_at_cursor(fmt='iso'))
            menu.addAction(insert_local)
            menu.addAction(insert_iso)

        if rpt_cfg.get('convert_timestamp', True):
            convert_act = QAction('Convert Selected Timestamp', self)
            convert_act.setShortcut(QKeySequence('Ctrl+Shift+C'))
            convert_act.triggered.connect(self.convert_selected_timestamp)
            menu.addAction(convert_act)

        if rpt_cfg.get('insert_template', True):
            tpl_act = QAction('Insert Template', self)
            tpl_act.triggered.connect(self.insert_template_from_context)
            menu.addAction(tpl_act)

        if rpt_cfg.get('validate_section', True):
            val_act = QAction('Validate Section', self)
            val_act.triggered.connect(self.validate_section)
            menu.addAction(val_act)

        if rpt_cfg.get('export_pdf', True):
            exp_act = QAction('Export Selection as PDF', self)
            exp_act.triggered.connect(self.export_selection_pdf)
            menu.addAction(exp_act)

        if rpt_cfg.get('embed_evidence', True):
            emb_act = QAction('Embed Evidence Reference', self)
            emb_act.triggered.connect(self.embed_evidence_reference)
            menu.addAction(emb_act)

        menu.exec_(self.report_editor.mapToGlobal(pos))

    # ---------------------- Report context menu stubs ----------------------
    def insert_template_from_context(self):
        try:
            tm = TemplateManager()
            tpl, ok = QInputDialog.getItem(self, 'Insert Template', 'Choose template:', tm.list_templates(), 0, False)
            if ok and tpl:
                content = tm.render(tpl, {})
                cursor = self.report_editor.textCursor()
                cursor.insertHtml(content)
                try:
                    self.audit.log('REPORT_TEMPLATE_INSERTED', {'template': tpl})
                except Exception as e:
                    logger.warning(f"Failed to log template insertion: {e}")
        except Exception as e:
            logger.error(f"Template insertion failed: {e}")
            QMessageBox.information(self, 'Insert Template', 'Template insertion not available.')

    def validate_section(self):
        cursor = self.report_editor.textCursor()
        text = cursor.selectedText() if cursor.hasSelection() else self.report_editor.toPlainText()
        # Very lightweight check: ensure presence of case number
        if 'Case' in text or self.case_data.get('case_number'):
            QMessageBox.information(self, 'Validate', 'Basic validation passed.')
        else:
            QMessageBox.warning(self, 'Validate', 'Validation: case number or required section missing.')

    def export_selection_pdf(self):
        cursor = self.report_editor.textCursor()
        content = cursor.selectedText() if cursor.hasSelection() else self.report_editor.toHtml()
        path, _ = QFileDialog.getSaveFileName(self, 'Export Selection as HTML', '', 'HTML Files (*.html)')
        if path:
            with open(path, 'w', encoding='utf-8') as f:
                f.write(content)
            try:
                self.audit.log('REPORT_SELECTION_EXPORTED', {'path': path})
            except Exception as e:
                logger.warning(f"Failed to log selection export: {e}")
            QMessageBox.information(self, 'Export', f'Section exported to {path}.')

    def embed_evidence_reference(self):
        eid, ok = QInputDialog.getText(self, 'Embed Evidence', 'Evidence ID or filename:')
        if ok and eid:
            cursor = self.report_editor.textCursor()
            cursor.insertHtml(f"<a href='#evidence:{eid}'>Evidence: {eid}</a>")
            try:
                self.audit.log('REPORT_EMBED_EVIDENCE', {'evidence': eid})
            except Exception as e:
                logger.warning(f"Failed to log evidence embed: {e}")

    def refresh_context_menu_settings(self):
        """Reload context-menu related settings from config and update toolbar/actions."""
        try:
            from auth import load_config
            cfg = load_config()
        except Exception as e:
            logger.warning(f"Could not load config for context menu refresh: {e}")
            cfg = {}
        rpt_cfg = cfg.get('context_menu', {}).get('reports', {}) if cfg else {}
        try:
            enabled = bool(rpt_cfg.get('insert_timestamp', True))
            # Report editor uses dynamic actions; no persistent toolbar actions to toggle in simple scaffold
        except Exception as e:
            logger.warning(f"Failed to get context menu settings: {e}")

    # ---------------------- Timestamp Conversion Helpers (report editor) ----------------------
    def _detect_tz_from_abbrev(self, text):
        if not text:
            return None
        import re
        m = re.search(r"\b([A-Z]{2,4}|UTC|GMT)\b", text)
        if not m:
            return None
        abb = m.group(1).upper()
        mapping = {
            'PST': ['America/Los_Angeles'], 'PDT': ['America/Los_Angeles'],
            'MST': ['America/Denver'], 'MDT': ['America/Denver'],
            'CST': ['America/Chicago', 'Asia/Shanghai'], 'CDT': ['America/Chicago'],
            'EST': ['America/New_York'], 'EDT': ['America/New_York'],
            'AKST': ['America/Anchorage'], 'AKDT': ['America/Anchorage'],
            'HST': ['Pacific/Honolulu'], 'UTC': ['UTC'], 'GMT': ['Etc/GMT']
        }
        vals = mapping.get(abb)
        if not vals:
            return None
        return vals[0]

    def _detect_tz_choices(self, text):
        if not text:
            return None
        import re
        m = re.search(r"\b([A-Z]{2,4}|UTC|GMT)\b", text)
        if not m:
            return None
        abb = m.group(1).upper()
        mapping = {
            'PST': ['America/Los_Angeles'], 'PDT': ['America/Los_Angeles'],
            'MST': ['America/Denver'], 'MDT': ['America/Denver'],
            'CST': ['America/Chicago', 'Asia/Shanghai'], 'CDT': ['America/Chicago'],
            'EST': ['America/New_York'], 'EDT': ['America/New_York'],
            'AKST': ['America/Anchorage'], 'AKDT': ['America/Anchorage'],
            'HST': ['Pacific/Honolulu'], 'UTC': ['UTC'], 'GMT': ['Etc/GMT']
        }
        return mapping.get(abb)

    def _parse_timestamp(self, text):
        if not text or not text.strip():
            return None, None, None
        try:
            if dateutil_parser:
                dt = dateutil_parser.parse(text, fuzzy=True)
                tzname = None
                if dt.tzinfo is None:
                    tzname = self._detect_tz_from_abbrev(text)
                    if tzname:
                        if ZoneInfo:
                            dt = dt.replace(tzinfo=ZoneInfo(tzname))
                        else:
                            dt = dt.replace(tzinfo=dateutil_tz.gettz(tzname))
                else:
                    # capture tz name if available
                    try:
                        tzname = dt.tzname()
                    except Exception as e:
                        logger.debug(f"Failed to get tzname: {e}")
                        tzname = None
                return dt, tzname, 'dateutil'
        except Exception as e:
            logger.debug(f"Dateutil parse failed: {e}")
        try:
            dt = datetime.fromisoformat(text)
            if dt.tzinfo is None:
                tzname = self._detect_tz_from_abbrev(text)
                if tzname:
                    if ZoneInfo:
                        dt = dt.replace(tzinfo=ZoneInfo(tzname))
                    else:
                        dt = dt.replace(tzinfo=dateutil_tz.gettz(tzname) if dateutil_tz else None)
            return dt, tzname, 'iso'
        except Exception as e:
            logger.debug(f"ISO parse failed: {e}")
        return None, None, None

    def insert_timestamp_at_cursor(self, fmt='local'):
        try:
            if fmt == 'iso':
                ts = datetime.now().astimezone().isoformat()
            else:
                ts = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            cursor = self.report_editor.textCursor()
            cursor.insertText(ts)
            self.report_editor.setFocus()
        except Exception as e:
            logger.warning(f"Failed to insert timestamp: {e}")

    def convert_selected_timestamp(self):
        cursor = self.report_editor.textCursor()
        if not cursor.hasSelection():
            QMessageBox.information(self, "Convert Timestamp", "Please highlight a timestamp to convert.")
            return
        selected = cursor.selectedText()
        dt, detected_tz, used = self._parse_timestamp(selected)
        if dt is None:
            QMessageBox.warning(self, "Convert Timestamp", "Could not parse the selected text as a timestamp.")
            return
        dlg = QDialog(self)
        dlg.setWindowTitle("Convert Timestamp")
        fl = QFormLayout(dlg)
        src_lbl = QLabel(selected)
        fl.addRow("Original:", src_lbl)
        src_tz_lbl = QLabel(detected_tz or "(not detected)")
        fl.addRow("Detected TZ:", src_tz_lbl)
        tz_combo = QComboBox()
        cfg_tz = None
        try:
            from auth import load_config
            cfg = load_config()
            cfg_tz = cfg.get('timezone')
        except Exception as e:
            logger.warning(f"Could not load config for timezone: {e}")
            cfg_tz = None
        choices = ['UTC', 'System Local']
        if cfg_tz and cfg_tz not in ['local', 'UTC']:
            choices.insert(0, cfg_tz)
        for z in ['America/New_York', 'America/Chicago', 'America/Denver', 'America/Los_Angeles']:
            if z not in choices:
                choices.append(z)
        tz_combo.addItems(choices)
        # Pre-select user's configured timezone when possible
        try:
            if cfg_tz:
                if cfg_tz == 'local':
                    tz_combo.setCurrentText('System Local')
                elif cfg_tz == 'UTC' or (isinstance(cfg_tz, str) and cfg_tz.startswith('UTC')):
                    tz_combo.setCurrentText('UTC')
                else:
                    tz_combo.setCurrentText(cfg_tz)
        except (ValueError, RuntimeError) as e:
            logger.debug(f"Could not set target TZ combo: {e}")
        fl.addRow("Target TZ:", tz_combo)
        fmt_combo = QComboBox()
        fmt_combo.addItems(['ISO 8601 (with TZ)', 'Readable (YYYY-MM-DD HH:MM:SS TZ)'])
        fl.addRow('Output Format:', fmt_combo)
        btns = QDialogButtonBox(QDialogButtonBox.Ok | QDialogButtonBox.Cancel)
        fl.addRow(btns)
        btns.accepted.connect(dlg.accept)
        btns.rejected.connect(dlg.reject)
        if dlg.exec_() != QDialog.Accepted:
            return
        target_tz_text = tz_combo.currentText()
        if target_tz_text == 'System Local':
            target_tz = None
        else:
            target_tz = target_tz_text
        try:
            if target_tz is None:
                if dateutil_tz:
                    tgt = dateutil_tz.tzlocal()
                else:
                    tgt = None
            else:
                if ZoneInfo:
                    tgt = ZoneInfo(target_tz)
                else:
                    tgt = dateutil_tz.gettz(target_tz) if dateutil_tz else None
            if dt.tzinfo is None:
                if detected_tz:
                    if ZoneInfo and isinstance(detected_tz, str):
                        dt = dt.replace(tzinfo=ZoneInfo(detected_tz))
                    elif dateutil_tz:
                        dt = dt.replace(tzinfo=dateutil_tz.gettz(detected_tz))
                else:
                    if cfg_tz and cfg_tz not in ['local', 'UTC']:
                        if ZoneInfo:
                            dt = dt.replace(tzinfo=ZoneInfo(cfg_tz))
                        elif dateutil_tz:
                            dt = dt.replace(tzinfo=dateutil_tz.gettz(cfg_tz))
                    elif cfg_tz == 'UTC':
                        if ZoneInfo:
                            dt = dt.replace(tzinfo=ZoneInfo('UTC'))
                        else:
                            dt = dt.replace(tzinfo=dateutil_tz.gettz('UTC') if dateutil_tz else None)
            if tgt is not None:
                converted = dt.astimezone(tgt)
            else:
                if dateutil_tz:
                    converted = dt.astimezone(dateutil_tz.tzlocal())
                else:
                    converted = dt
            if fmt_combo.currentText().startswith('ISO'):
                out = converted.isoformat()
            else:
                try:
                    tzname = converted.tzname() or ''
                except Exception:
                    tzname = ''
                out = converted.strftime('%Y-%m-%d %H:%M:%S') + (f' {tzname}' if tzname else '')
            cursor.insertText(out)
            # Log replacement
            try:
                self.audit.log('REPORT_TIMESTAMP_CONVERTED', {'original': selected, 'converted': out, 'case': self.case_data.get('case_number')})
            except Exception as e:
                logger.warning(f"Failed to log timestamp conversion: {e}")
        except Exception as e:
            QMessageBox.warning(self, 'Convert Timestamp', f'Error converting timestamp: {e}')
        

    def toggle_bold(self, checked):
        cursor = self.report_editor.textCursor()
        if cursor.hasSelection():
            fmt = cursor.charFormat()
            fmt.setFontWeight(QFont.Bold if checked else QFont.Normal)
            cursor.setCharFormat(fmt)
        fmt = self.report_editor.currentCharFormat()
        fmt.setFontWeight(QFont.Bold if checked else QFont.Normal)
        self.report_editor.setCurrentCharFormat(fmt)

    def toggle_italic(self):
        cursor = self.report_editor.textCursor()
        if cursor.hasSelection():
            fmt = cursor.charFormat()
            fmt.setFontItalic(not fmt.fontItalic())
            cursor.setCharFormat(fmt)
        else:
            fmt = self.report_editor.currentCharFormat()
            fmt.setFontItalic(not fmt.fontItalic())
            self.report_editor.setCurrentCharFormat(fmt)

    def toggle_underline(self):
        cursor = self.report_editor.textCursor()
        if cursor.hasSelection():
            fmt = cursor.charFormat()
            fmt.setFontUnderline(not fmt.fontUnderline())
            cursor.setCharFormat(fmt)
        else:
            fmt = self.report_editor.currentCharFormat()
            fmt.setFontUnderline(not fmt.fontUnderline())
            self.report_editor.setCurrentCharFormat(fmt)

    def insert_bullet_list(self):
        cursor = self.report_editor.textCursor()
        list_format = QTextListFormat()
        list_format.setStyle(QTextListFormat.ListDisc)
        cursor.createList(list_format)

    def insert_numbered_list(self):
        cursor = self.report_editor.textCursor()
        list_format = QTextListFormat()
        list_format.setStyle(QTextListFormat.ListDecimal)
        cursor.createList(list_format)

    def export_pdf(self, finalize=False):
        try:
            html_template = f"""<html><head><meta charset="utf-8"></head><body>{self.report_editor.toHtml()}</body></html>"""

            # Ask user where to save the PDF
            pdf_path, _ = QFileDialog.getSaveFileName(self, "Save PDF", os.path.join(self.case_dir, "report.pdf"), "PDF Files (*.pdf)")
            if not pdf_path:
                return

            try:
                from weasyprint import HTML
                import warnings
                with warnings.catch_warnings():
                    warnings.simplefilter("ignore")
                    HTML(string=html_template).write_pdf(pdf_path, presentational_hints=True)

                # Compute hash with progress callback for large files
                def progress_callback(bytes_processed, total_bytes):
                    if total_bytes > 10 * 1024 * 1024:  # Show progress for files > 10MB
                        progress = int((bytes_processed / total_bytes) * 100)
                        self.hash_label.setText(f"Computing hash... {progress}%")

                pdf_hash = compute_sha256(pdf_path, progress_callback)
                self.hash_label.setText(f"Final PDF Hash (SHA-256):\n{pdf_hash}")
                self.saved_pdf_hash = pdf_hash

                if finalize:
                    try:
                        self.db.save_report(self.case_data, self.report_editor.toHtml(), self.appendices, pdf_hash)
                    except Exception as e:
                        logger.warning(f"Failed to save report after PDF export: {e}")

                try:
                    self.audit.log_pdf_finalized(pdf_path, pdf_hash)
                except Exception as e:
                    logger.warning(f"Failed to log PDF finalization: {e}")

                QMessageBox.information(self, "Success", f"PDF saved: {pdf_path}\nHash: {pdf_hash}")
            except ImportError:
                QMessageBox.warning(self, "PDF Export Disabled", "WeasyPrint is not available. PDF export features are disabled.")
            except Exception as e:
                QMessageBox.critical(self, "Error", f"PDF generation failed:\n{str(e)}")
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Export failed: {e}")

    def add_appendix(self):
        file_path, _ = QFileDialog.getOpenFileName(self, "Add Appendix", "", "All Files (*)")
        if file_path:
            dest_path = os.path.join(self.case_dir, "appendices", os.path.basename(file_path))
            os.makedirs(os.path.dirname(dest_path), exist_ok=True)
            shutil.copy(file_path, dest_path)
            self.appendices.append(dest_path)
            self.update_appendices_list()
            self.audit.log("APPENDIX_ADDED", {"file": os.path.basename(dest_path)})

    def view_appendices(self):
        dialog = QDialog(self)
        dialog.setWindowTitle("View Appendices")
        layout = QVBoxLayout(dialog)
        appendices_list = QListWidget()
        for appendix in self.appendices:
            appendices_list.addItem(os.path.basename(appendix))
        layout.addWidget(appendices_list)
        buttons = QDialogButtonBox(QDialogButtonBox.Ok)
        buttons.accepted.connect(dialog.accept)
        layout.addWidget(buttons)
        dialog.exec_()

    def remove_appendix(self):
        if not self.appendices:
            QMessageBox.information(self, "No Appendices", "There are no appendices to remove.")
            return
        dialog = QDialog(self)
        dialog.setWindowTitle("Remove Appendix")
        layout = QVBoxLayout(dialog)
        appendices_list = QListWidget()
        for appendix in self.appendices:
            appendices_list.addItem(os.path.basename(appendix))
        layout.addWidget(appendices_list)
        buttons = QDialogButtonBox(QDialogButtonBox.Ok | QDialogButtonBox.Cancel)
        buttons.accepted.connect(dialog.accept)
        buttons.rejected.connect(dialog.reject)
        layout.addWidget(buttons)
        if dialog.exec_() == QDialog.Accepted:
            selected_items = appendices_list.selectedItems()
            if selected_items:
                selected_item = selected_items[0]
                index = appendices_list.row(selected_item)
                removed_path = self.appendices.pop(index)
                self.audit.log("APPENDIX_REMOVED", {"file": os.path.basename(removed_path)})
                QMessageBox.information(self, "Removed", f"Appendix '{os.path.basename(removed_path)}' removed.")
            else:
                QMessageBox.warning(self, "No Selection", "Please select an appendix to remove.")

    def update_appendices_list(self):
        # Placeholder method to update appendices list if needed
        pass


def validate_appendix_path(case_dir, file_path):
    """
    Validate that an appendix file path is within the case directory.
    Prevents path traversal attacks (e.g., ../../sensitive_file.txt)
    
    Args:
        case_dir: Base case directory path
        file_path: File path to validate
        
    Returns:
        str: The validated file path
        
    Raises:
        ValueError: If path is outside case directory
    """
    # Get absolute paths to prevent traversal
    real_case_dir = os.path.abspath(case_dir)
    real_file_path = os.path.abspath(file_path)
    
    # Check if file is within case directory
    if not real_file_path.startswith(real_case_dir + os.sep) and real_file_path != real_case_dir:
        raise ValueError(f"Path traversal detected: {file_path} is outside case directory {case_dir}")
    
    return real_file_path


class ReportsWindow(QMainWindow):
    def __init__(self, case_data, db_manager, audit_logger, current_user):
        try:
            super().__init__()
            self.setWindowTitle(f"Reports - Case {case_data['case_number']}")
            self.setMinimumSize(1200, 800)

            self.reports_tab = ReportsTab(case_data, db_manager, audit_logger, current_user, self)
            self.setCentralWidget(self.reports_tab)

            self.setup_menu()
            self.setup_toolbar()
        except Exception as e:
            logger.error(f"Error initializing ReportsWindow: {e}")
            raise

    def setup_menu(self):
        menu = self.menuBar()

        file_menu = menu.addMenu("File")
        save_act = QAction("Save", self)
        save_act.setShortcut(QKeySequence.Save)
        save_act.triggered.connect(self.save_report)
        file_menu.addAction(save_act)

        export_act = QAction("Export PDF", self)
        export_act.setShortcut(QKeySequence("Ctrl+E"))
        export_act.triggered.connect(self.reports_tab.export_pdf)
        file_menu.addAction(export_act)

        finalize_act = QAction("Finalize Report", self)
        finalize_act.triggered.connect(lambda: self.reports_tab.export_pdf(finalize=True))
        file_menu.addAction(finalize_act)

        export_peer_review_act = QAction("Export for Peer Review", self)
        export_peer_review_act.triggered.connect(self.export_for_peer_review)
        file_menu.addAction(export_peer_review_act)

        import_peer_review_act = QAction("Import Peer Review", self)
        import_peer_review_act.triggered.connect(self.import_peer_review)
        file_menu.addAction(import_peer_review_act)

        load_nist_template_act = QAction("Load NIST Template", self)
        load_nist_template_act.triggered.connect(self.load_nist_template)
        file_menu.addAction(load_nist_template_act)

        import_template_act = QAction("Import Template from DOCX", self)
        import_template_act.triggered.connect(self.import_template_from_docx)
        file_menu.addAction(import_template_act)

        close_act = QAction("Close", self)
        close_act.setShortcut(QKeySequence.Close)
        close_act.triggered.connect(self.close)
        file_menu.addAction(close_act)

        edit_menu = menu.addMenu("Edit")
        undo_act = QAction("Undo", self)
        undo_act.setShortcut(QKeySequence.Undo)
        undo_act.triggered.connect(self.reports_tab.report_editor.undo)
        edit_menu.addAction(undo_act)

        redo_act = QAction("Redo", self)
        redo_act.setShortcut(QKeySequence.Redo)
        redo_act.triggered.connect(self.reports_tab.report_editor.redo)
        edit_menu.addAction(redo_act)

        find_replace_act = QAction("Find and Replace", self)
        find_replace_act.setShortcut(QKeySequence.Find)
        find_replace_act.triggered.connect(self.reports_tab.word_processor.show_find_replace)
        edit_menu.addAction(find_replace_act)

        insert_menu = menu.addMenu("Insert")
        image_act = QAction("Image", self)
        image_act.triggered.connect(self.insert_image)
        insert_menu.addAction(image_act)

        table_act = QAction("Table", self)
        table_act.triggered.connect(self.insert_table)
        insert_menu.addAction(table_act)

        page_break_act = QAction("Page Break", self)
        page_break_act.triggered.connect(self.insert_page_break)
        insert_menu.addAction(page_break_act)

        note_from_notes_act = QAction("Note from Notes Tab", self)
        note_from_notes_act.triggered.connect(self.insert_note_from_notes)
        insert_menu.addAction(note_from_notes_act)

        format_menu = menu.addMenu("Format")
        font_act = QAction("Font...", self)
        font_act.triggered.connect(self.change_font)
        format_menu.addAction(font_act)

        color_act = QAction("Text Color...", self)
        color_act.triggered.connect(self.change_text_color)
        format_menu.addAction(color_act)

        bg_color_act = QAction("Background Color...", self)
        bg_color_act.triggered.connect(self.change_bg_color)
        format_menu.addAction(bg_color_act)

        align_left_act = QAction("Align Left", self)
        align_left_act.triggered.connect(lambda: self.align_text(Qt.AlignLeft))
        format_menu.addAction(align_left_act)

        align_center_act = QAction("Align Center", self)
        align_center_act.triggered.connect(lambda: self.align_text(Qt.AlignCenter))
        format_menu.addAction(align_center_act)

        align_right_act = QAction("Align Right", self)
        align_right_act.triggered.connect(lambda: self.align_text(Qt.AlignRight))
        format_menu.addAction(align_right_act)

        align_justify_act = QAction("Justify", self)
        align_justify_act.triggered.connect(lambda: self.align_text(Qt.AlignJustify))
        format_menu.addAction(align_justify_act)

        appendices_menu = menu.addMenu("Appendices")
        view_appendices_act = QAction("View Appendices", self)
        view_appendices_act.triggered.connect(self.reports_tab.view_appendices)
        appendices_menu.addAction(view_appendices_act)

        add_appendix_act = QAction("Add Appendix", self)
        add_appendix_act.triggered.connect(self.reports_tab.add_appendix)
        appendices_menu.addAction(add_appendix_act)

        remove_appendix_act = QAction("Remove Appendix", self)
        remove_appendix_act.triggered.connect(self.reports_tab.remove_appendix)
        appendices_menu.addAction(remove_appendix_act)

    def setup_toolbar(self):
        toolbar = self.addToolBar("Formatting")

        # Font family
        self.font_combo = QFontComboBox()
        self.font_combo.currentFontChanged.connect(self.change_font_family)
        toolbar.addWidget(self.font_combo)

        # Font size
        self.size_combo = QComboBox()
        self.size_combo.addItems([str(i) for i in range(8, 72, 2)])
        self.size_combo.setCurrentText("12")
        self.size_combo.currentTextChanged.connect(self.change_font_size)
        toolbar.addWidget(self.size_combo)

        toolbar.addSeparator()

        # Style dropdown
        style_btn = QToolButton(self)
        style_btn.setText("Style")
        style_btn.setPopupMode(QToolButton.InstantPopup)
        style_menu = QMenu(self)
        style_btn.setMenu(style_menu)

        bold_act = QAction("Bold", self)
        bold_act.setCheckable(True)
        bold_act.triggered.connect(self.reports_tab.toggle_bold)
        style_menu.addAction(bold_act)

        italic_act = QAction("Italic", self)
        italic_act.setCheckable(True)
        italic_act.triggered.connect(self.reports_tab.toggle_italic)
        style_menu.addAction(italic_act)

        underline_act = QAction("Underline", self)
        underline_act.setCheckable(True)
        underline_act.triggered.connect(self.reports_tab.toggle_underline)
        style_menu.addAction(underline_act)

        strikethrough_act = QAction("Strikethrough", self)
        strikethrough_act.triggered.connect(self.reports_tab.word_processor.apply_strikethrough)
        style_menu.addAction(strikethrough_act)

        style_menu.addSeparator()

        subscript_act = QAction("Subscript", self)
        subscript_act.triggered.connect(self.reports_tab.word_processor.apply_subscript)
        style_menu.addAction(subscript_act)

        superscript_act = QAction("Superscript", self)
        superscript_act.triggered.connect(self.reports_tab.word_processor.apply_superscript)
        style_menu.addAction(superscript_act)

        toolbar.addWidget(style_btn)

        toolbar.addSeparator()

        # Color dropdown
        color_btn = QToolButton(self)
        color_btn.setText("Color")
        color_btn.setPopupMode(QToolButton.InstantPopup)
        color_menu = QMenu(self)
        color_btn.setMenu(color_menu)

        color_act = QAction("Text Color", self)
        color_act.triggered.connect(self.change_text_color)
        color_menu.addAction(color_act)

        bg_color_act = QAction("Background Color", self)
        bg_color_act.triggered.connect(self.change_bg_color)
        color_menu.addAction(bg_color_act)

        toolbar.addWidget(color_btn)

        toolbar.addSeparator()

        # Alignment dropdown
        align_btn = QToolButton(self)
        align_btn.setText("Align")
        align_btn.setPopupMode(QToolButton.InstantPopup)
        align_menu = QMenu(self)
        align_btn.setMenu(align_menu)

        align_left_act = QAction("Align Left", self)
        align_left_act.triggered.connect(lambda: self.align_text(Qt.AlignLeft))
        align_menu.addAction(align_left_act)

        align_center_act = QAction("Align Center", self)
        align_center_act.triggered.connect(lambda: self.align_text(Qt.AlignCenter))
        align_menu.addAction(align_center_act)

        align_right_act = QAction("Align Right", self)
        align_right_act.triggered.connect(lambda: self.align_text(Qt.AlignRight))
        align_menu.addAction(align_right_act)

        align_justify_act = QAction("Justify", self)
        align_justify_act.triggered.connect(lambda: self.align_text(Qt.AlignJustify))
        align_menu.addAction(align_justify_act)

        toolbar.addWidget(align_btn)

        toolbar.addSeparator()

        # Lists dropdown
        list_btn = QToolButton(self)
        list_btn.setText("Lists")
        list_btn.setPopupMode(QToolButton.InstantPopup)
        list_menu = QMenu(self)
        list_btn.setMenu(list_menu)

        bullet_act = QAction("Bullet List", self)
        bullet_act.triggered.connect(self.reports_tab.insert_bullet_list)
        list_menu.addAction(bullet_act)

        number_act = QAction("Numbered List", self)
        number_act.triggered.connect(self.reports_tab.insert_numbered_list)
        list_menu.addAction(number_act)

        toolbar.addWidget(list_btn)

        toolbar.addSeparator()

        # Indent dropdown
        indent_btn = QToolButton(self)
        indent_btn.setText("Indent")
        indent_btn.setPopupMode(QToolButton.InstantPopup)
        indent_menu = QMenu(self)
        indent_btn.setMenu(indent_menu)

        indent_act = QAction("Increase Indent", self)
        indent_act.triggered.connect(self.increase_indent)
        indent_menu.addAction(indent_act)

        outdent_act = QAction("Decrease Indent", self)
        outdent_act.triggered.connect(self.decrease_indent)
        indent_menu.addAction(outdent_act)

        toolbar.addWidget(indent_btn)

        # Timestamp tools (insert/convert)
        ts_btn = QToolButton(self)
        ts_btn.setText('Timestamps')
        ts_btn.setPopupMode(QToolButton.InstantPopup)
        ts_menu = QMenu(self)

        insert_local = QAction('Insert Date/Time (Local)', self)
        insert_local.triggered.connect(lambda: self.reports_tab.insert_timestamp_at_cursor(fmt='local'))
        ts_menu.addAction(insert_local)

        insert_iso = QAction('Insert ISO Date/Time (with TZ)', self)
        insert_iso.triggered.connect(lambda: self.reports_tab.insert_timestamp_at_cursor(fmt='iso'))
        ts_menu.addAction(insert_iso)

        convert_act = QAction('Convert Selected Timestamp', self)
        convert_act.triggered.connect(self.reports_tab.convert_selected_timestamp)
        ts_menu.addAction(convert_act)

        ts_btn.setMenu(ts_menu)
        toolbar.addWidget(ts_btn)

    def save_report(self):
        self.reports_tab.db.save_report(self.reports_tab.case_data, self.reports_tab.report_editor.toHtml(), self.reports_tab.appendices, self.reports_tab.saved_pdf_hash)
        self.reports_tab.audit.log("REPORT_SAVED", {"case": self.reports_tab.case_data['case_number']})

    def change_font_family(self, font):
        cursor = self.reports_tab.report_editor.textCursor()
        if cursor.hasSelection():
            fmt = cursor.charFormat()
            fmt.setFontFamily(font.family())
            cursor.setCharFormat(fmt)
        else:
            fmt = self.reports_tab.report_editor.currentCharFormat()
            fmt.setFontFamily(font.family())
            self.reports_tab.report_editor.setCurrentCharFormat(fmt)

    def change_font_size(self, size):
        cursor = self.reports_tab.report_editor.textCursor()
        if cursor.hasSelection():
            fmt = cursor.charFormat()
            fmt.setFontPointSize(int(size))
            cursor.setCharFormat(fmt)
        else:
            fmt = self.reports_tab.report_editor.currentCharFormat()
            fmt.setFontPointSize(int(size))
            self.reports_tab.report_editor.setCurrentCharFormat(fmt)

    def change_font(self):
        from PyQt5.QtWidgets import QFontDialog
        font, ok = QFontDialog.getFont()
        if ok:
            cursor = self.reports_tab.report_editor.textCursor()
            if cursor.hasSelection():
                fmt = cursor.charFormat()
                fmt.setFont(font)
                cursor.setCharFormat(fmt)
            else:
                fmt = self.reports_tab.report_editor.currentCharFormat()
                fmt.setFont(font)
                self.reports_tab.report_editor.setCurrentCharFormat(fmt)

    def change_text_color(self):
        color = QColorDialog.getColor()
        if color.isValid():
            cursor = self.reports_tab.report_editor.textCursor()
            if cursor.hasSelection():
                fmt = cursor.charFormat()
                fmt.setForeground(color)
                cursor.setCharFormat(fmt)
            else:
                fmt = self.reports_tab.report_editor.currentCharFormat()
                fmt.setForeground(color)
                self.reports_tab.report_editor.setCurrentCharFormat(fmt)

    def change_bg_color(self):
        color = QColorDialog.getColor()
        if color.isValid():
            cursor = self.reports_tab.report_editor.textCursor()
            if cursor.hasSelection():
                fmt = cursor.charFormat()
                fmt.setBackground(color)
                cursor.setCharFormat(fmt)
            else:
                fmt = self.reports_tab.report_editor.currentCharFormat()
                fmt.setBackground(color)
                self.reports_tab.report_editor.setCurrentCharFormat(fmt)

    def align_text(self, alignment):
        cursor = self.reports_tab.report_editor.textCursor()
        block_fmt = cursor.blockFormat()
        block_fmt.setAlignment(alignment)
        cursor.setBlockFormat(block_fmt)

    def increase_indent(self):
        cursor = self.reports_tab.report_editor.textCursor()
        block_fmt = cursor.blockFormat()
        block_fmt.setIndent(block_fmt.indent() + 1)
        cursor.setBlockFormat(block_fmt)

    def decrease_indent(self):
        cursor = self.reports_tab.report_editor.textCursor()
        block_fmt = cursor.blockFormat()
        indent = max(0, block_fmt.indent() - 1)
        block_fmt.setIndent(indent)
        cursor.setBlockFormat(block_fmt)

    def insert_image(self):
        file, _ = QFileDialog.getOpenFileName(self, "Insert Image", "", "Images (*.png *.jpg *.jpeg *.bmp *.gif)")
        if file:
            cursor = self.reports_tab.report_editor.textCursor()
            cursor.insertHtml(f'<img src="{file}" />')

    def insert_table(self):
        self.reports_tab.word_processor.insert_advanced_table()

    def insert_page_break(self):
        cursor = self.reports_tab.report_editor.textCursor()
        cursor.insertHtml('<div style="page-break-before: always;"></div>')

    def insert_note_from_notes(self):
        """Insert a note from the notes tab into the report editor at cursor position"""
        notes_dir = os.path.join("cases", self.reports_tab.case_data['case_number'], "notes")
        notes_file = os.path.join(notes_dir, "notes_data.json")

        if not os.path.exists(notes_file):
            QMessageBox.warning(self, "No Notes", "No notes found for this case.")
            return

        try:
            with open(notes_file, 'r') as f:
                notes_data = json.load(f)
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to load notes: {str(e)}")
            return

        notes = notes_data.get('notes', {})
        if not notes:
            QMessageBox.information(self, "No Notes", "No notes available to insert.")
            return

        # Create dialog to select note
        dialog = QDialog(self)
        dialog.setWindowTitle("Select Note to Insert")
        layout = QVBoxLayout(dialog)

        notes_list = QListWidget()
        for note_id, note in notes.items():
            name = note.get('name', f"Note {note_id}")
            timestamp = note.get('timestamp', '')
            if timestamp:
                try:
                    dt = datetime.fromisoformat(timestamp)
                    timestamp_str = dt.strftime('%Y-%m-%d %H:%M')
                except ValueError as e:
                    timestamp_str = timestamp
            else:
                timestamp_str = 'Unknown'
            item_text = f"{name} ({timestamp_str})"
            item = QListWidgetItem(item_text)
            item.setData(Qt.UserRole, note_id)
            notes_list.addItem(item)

        layout.addWidget(notes_list)

        buttons = QDialogButtonBox(QDialogButtonBox.Ok | QDialogButtonBox.Cancel)
        buttons.accepted.connect(dialog.accept)
        buttons.rejected.connect(dialog.reject)
        layout.addWidget(buttons)

        if dialog.exec_() == QDialog.Accepted:
            selected_items = notes_list.selectedItems()
            if selected_items:
                selected_item = selected_items[0]
                note_id = selected_item.data(Qt.UserRole)
                note_content = notes[note_id].get('content', '')

                if note_content:
                    # Insert at cursor position
                    cursor = self.reports_tab.report_editor.textCursor()
                    cursor.insertHtml(note_content)
                    self.reports_tab.audit.log("NOTE_INSERTED_INTO_REPORT", {"note_id": note_id, "case": self.reports_tab.case_data['case_number']})
                    QMessageBox.information(self, "Note Inserted", "Note has been inserted into the report.")
                else:
                    QMessageBox.warning(self, "Empty Note", "The selected note is empty.")
            else:
                QMessageBox.warning(self, "No Selection", "Please select a note to insert.")

    def export_for_peer_review(self):
        """Export report data for peer review"""
        from PyQt5.QtWidgets import QFileDialog
        import json

        # Prepare report data for export
        report_data = {
            "case_info": {
                "case_number": self.reports_tab.case_data.get('case_number', ''),
                "title": self.reports_tab.case_data.get('title', 'Untitled Case')
            },
            "report_html": self.reports_tab.report_editor.toHtml(),
            "appendices": self.reports_tab.appendices,
            "pdf_hash": self.reports_tab.saved_pdf_hash,
            "exported_at": datetime.now().isoformat(),
            "exported_by": getattr(self.reports_tab, 'current_user', 'Unknown')
        }

        # Save to file
        file_path, _ = QFileDialog.getSaveFileName(
            self, "Export for Peer Review", "",
            "Peer Review Files (*.peerreview.json);;All Files (*)"
        )

        if file_path:
            try:
                with open(file_path, 'w', encoding='utf-8') as f:
                    json.dump(report_data, f, indent=2, ensure_ascii=False)

                QMessageBox.information(self, "Exported",
                    f"Report exported for peer review:\n{file_path}")

                self.reports_tab.audit.log("REPORT_EXPORTED_FOR_PEER_REVIEW",
                    {"file": file_path, "case": self.reports_tab.case_data['case_number']})

            except Exception as e:
                QMessageBox.critical(self, "Export Error",
                    f"Failed to export report:\n{str(e)}")

    def import_peer_review(self):
        """Import peer review data from a file"""
        from PyQt5.QtWidgets import QFileDialog
        import json

        # Open file dialog to select reviewed file
        file_path, _ = QFileDialog.getOpenFileName(
            self, "Import Peer Review", "",
            "Reviewed Report Files (*.reviewed.json);;All Files (*)"
        )

        if file_path:
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    review_data = json.load(f)

                # Validate the review data
                if 'review_data' not in review_data:
                    QMessageBox.warning(self, "Invalid File",
                        "The selected file does not contain valid peer review data.")
                    return

                review_info = review_data['review_data']

                # Display review summary
                summary = f"Peer Review Imported\n\n"
                summary += f"Reviewer: {review_info.get('reviewer_info', {}).get('name', 'Unknown')}\n"
                summary += f"Agency: {review_info.get('reviewer_info', {}).get('agency', 'Unknown')}\n"
                summary += f"Role: {review_info.get('reviewer_info', {}).get('role', 'Unknown')}\n\n"
                summary += f"Comments: {len(review_info.get('comments', []))}\n"
                summary += f"Summary:\n{review_info.get('summary', 'No summary provided')}"

                QMessageBox.information(self, "Peer Review Imported", summary)

                # Log the import
                self.reports_tab.audit.log("PEER_REVIEW_IMPORTED",
                    {"file": file_path, "case": self.reports_tab.case_data['case_number']})

            except Exception as e:
                QMessageBox.critical(self, "Import Error",
                    f"Failed to import peer review:\n{str(e)}")

    def load_nist_template(self):
        """Load the prebuilt NIST template into the report editor"""
        from templates import DEFAULT_TEMPLATES

        nist_template = DEFAULT_TEMPLATES.get("SWGDE/NIST Standard", "")
        if nist_template:
            # Replace placeholders with case data
            case = self.reports_tab.case_data
            case_number = case.get('case_number', '')

            # Get evidence details for the case
            evidence_table_html = self._build_evidence_table(case_number)

            # Replace the evidence table placeholder in the template
            nist_template = nist_template.replace(
                '<table border="1" cellpadding="5">\n<tr><th>Item ID</th><th>Description</th><th>Make/Model</th><th>Serial/IMEI</th><th>Hash (Acquisition)</th></tr>\n<tr><td></td><td></td><td></td><td></td><td></td></tr>\n</table>',
                evidence_table_html
            )

            html = nist_template.format(
                case_number=case_number,
                suspect=case.get('suspect', 'N/A'),
                investigator=case.get('investigator', ''),
                agency=case.get('agency', ''),
                date=datetime.now().strftime('%B %d, %Y')
            )
            self.reports_tab.report_editor.setHtml(html)
            QMessageBox.information(self, "Template Loaded", "NIST template loaded successfully with evidence details.")
            self.reports_tab.audit.log("NIST_TEMPLATE_LOADED", {"case": case_number})
        else:
            QMessageBox.warning(self, "Template Error", "NIST template not found.")

    def _build_evidence_table(self, case_number):
        """Build HTML table with evidence details for the case"""
        try:
            # Import the evidence details function
            from cases_bp import get_evidence_details

            evidence_items = get_evidence_details(case_number)

            if not evidence_items:
                # Return empty table if no evidence
                return '<table border="1" cellpadding="5">\n<tr><th>Item ID</th><th>Description</th><th>Make/Model</th><th>Serial/IMEI</th><th>Hash (Acquisition)</th></tr>\n<tr><td></td><td></td><td></td><td></td><td></td></tr>\n</table>'

            # Build table rows
            table_html = '<table border="1" cellpadding="5">\n'
            table_html += '<tr><th>Item ID</th><th>Description</th><th>Make/Model</th><th>Serial/IMEI</th><th>Hash (Acquisition)</th></tr>\n'

            for item in evidence_items:
                item_id = str(item.get('id', ''))
                description = item.get('evidence_item_number', '') or item.get('type', '')
                make_model = item.get('make_model', '') or item.get('type', '')
                serial_imei = item.get('serial_number', '') or item.get('imei', '')
                hash_value = item.get('acquisition_hash', '') or item.get('hash', '')

                table_html += f'<tr><td>{item_id}</td><td>{description}</td><td>{make_model}</td><td>{serial_imei}</td><td>{hash_value}</td></tr>\n'

            # Add an empty row for manual editing
            table_html += '<tr><td></td><td></td><td></td><td></td><td></td></tr>\n'
            table_html += '</table>'

            return table_html

        except Exception as e:
            logger.error(f"Error building evidence table: {e}")
            # Return basic empty table on error
            return '<table border="1" cellpadding="5">\n<tr><th>Item ID</th><th>Description</th><th>Make/Model</th><th>Serial/IMEI</th><th>Hash (Acquisition)</th></tr>\n<tr><td></td><td></td><td></td><td></td><td></td></tr>\n</table>'

    def import_template_from_docx(self):
        """Import a template from a Microsoft Office DOCX file"""
        from PyQt5.QtWidgets import QFileDialog

        file_path, _ = QFileDialog.getOpenFileName(
            self, "Import Template from DOCX", "",
            "Word Documents (*.docx);;All Files (*)"
        )

        if file_path:
            try:
                # Try to import python-docx
                try:
                    from docx import Document
                    doc = Document(file_path)
                    html_content = ""

                    for paragraph in doc.paragraphs:
                        if paragraph.text.strip():
                            # Convert basic formatting to HTML
                            text = paragraph.text
                            if paragraph.runs:
                                for run in paragraph.runs:
                                    if run.bold:
                                        text = text.replace(run.text, f"<b>{run.text}</b>")
                                    if run.italic:
                                        text = text.replace(run.text, f"<i>{run.text}</i>")
                                    if run.underline:
                                        text = text.replace(run.text, f"<u>{run.text}</u>")
                            html_content += f"<p>{text}</p>"

                    # Handle tables
                    for table in doc.tables:
                        html_content += "<table border='1' style='border-collapse: collapse;'>"
                        for row in table.rows:
                            html_content += "<tr>"
                            for cell in row.cells:
                                html_content += f"<td style='padding: 5px;'>{cell.text}</td>"
                            html_content += "</tr>"
                        html_content += "</table>"

                    if html_content:
                        self.reports_tab.report_editor.setHtml(html_content)
                        QMessageBox.information(self, "Template Imported",
                            f"Template imported from:\n{file_path}")
                        self.reports_tab.audit.log("TEMPLATE_IMPORTED_FROM_DOCX",
                            {"file": file_path, "case": self.reports_tab.case_data.get('case_number', '')})
                    else:
                        QMessageBox.warning(self, "Import Warning",
                            "No content found in the selected DOCX file.")

                except ImportError:
                    QMessageBox.critical(self, "Missing Dependency",
                        "python-docx is required to import DOCX templates.\n\n"
                        "Install it with: pip install python-docx")

            except Exception as e:
                QMessageBox.critical(self, "Import Error",
                    f"Failed to import template:\n{str(e)}")

    def closeEvent(self, event):
        self.save_report()
        super().closeEvent(event)
